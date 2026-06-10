from typing import Optional
#!/usr/bin/env python3
"""
Параллельная генерация SEO-текстов через Gemini API.
Читает task.json, генерирует тексты, сохраняет DOCX, загружает в Dropbox.
"""
import asyncio, json, os, sys, datetime, zipfile, re, glob as _glob
import httpx

# Telegram bot config
TELEGRAM_BOT_TOKEN = ""  # Disabled — use web UI only  # Disabled — use web UI only
TELEGRAM_ALLOWED_USERS = ["200063243", "7687757525"]
TELEGRAM_GROUP_ID = "-5225581427"  # Группа с Софией

def tg_send(*args, **kwargs):  # TG disabled
    pass  # Disabled — use web UI
WORKSPACE = os.path.dirname(os.path.abspath(__file__))
SECRETS_PATH = os.path.join(WORKSPACE, "SECRETS.md")
TASK_PATH = os.environ.get("TASK_FILE_OVERRIDE") or os.path.join(WORKSPACE, "task.json")

def read_secret(key):
    """Читаем ключ из SECRETS.md"""
    with open(SECRETS_PATH, encoding="utf-8") as f:
        for line in f:
            if key in line and ":" in line:
                return line.split(":", 1)[1].strip()
    return None

# Load both Gemini keys from SECRETS.md
try:
    with open(SECRETS_PATH) as f:
        _secrets_content = f.read()
    _all_gemini_keys = re.findall(r'AIzaSy\S+', _secrets_content)
    GEMINI_KEY = _all_gemini_keys[0] if _all_gemini_keys else None
    GEMINI_KEY_2 = _all_gemini_keys[1] if len(_all_gemini_keys) > 1 else None
    _k1 = read_secret("Key: AIza")
    _k2 = read_secret("Key2: AIza") or read_secret("Gemini API Key 2")
    if _k1: GEMINI_KEY = _k1
    if _k2: GEMINI_KEY_2 = _k2
except Exception as _ke:
    GEMINI_KEY = None
    GEMINI_KEY_2 = None

# Key rotation
_key_rotation_idx = 0
def get_next_gemini_key():
    global _key_rotation_idx
    keys = [k for k in [GEMINI_KEY, GEMINI_KEY_2] if k]
    key = keys[_key_rotation_idx % len(keys)]
    _key_rotation_idx += 1
    return key

DROPBOX_APP_KEY = "wczm5uvf7mqvu5z"
DROPBOX_APP_SECRET = "6962f0jns0nnajb"
DROPBOX_REFRESH = "UFvU3pBeShIAAAAAAAAAAVHI1vIRhPdHePMWsanDK9bAGQM4zctvTcE7gcq3vqTy"

GEMINI_MODEL = "google/gemini-2.5-flash"  # default fallback
def _get_generation_model():
    """Читаем модель из task.json, fallback на GEMINI_MODEL."""
    import json, os
    task_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "task.json")
    try:
        with open(task_path) as f:
            t = json.load(f)
        m = t.get("model", "").strip()
        if m:
            return m
    except Exception:
        pass
    return GEMINI_MODEL

GEMINI_URL = "https://openrouter.ai/api/v1/chat/completions"
OPENROUTER_KEY = "YOUR_OPENROUTER_KEY_HERE"
OPENROUTER_KEY_2 = "YOUR_OPENROUTER_KEY_HERE"
# ─── v3: Ahrefs SERP Analysis ──────────────────────────────────────────────────
AHREFS_TOKEN = "2hHH9gkDW15WZSzdHvWh8jlWujY15glYymvIttvN"


_or_key_idx = 0
def get_next_or_key():
    global _or_key_idx
    keys = [k for k in [OPENROUTER_KEY, OPENROUTER_KEY_2] if k]
    key = keys[_or_key_idx % len(keys)]
    _or_key_idx += 1
    return key

# ─── Dropbox ──────────────────────────────────────────────────────────────────
def get_dropbox_token():
    import requests
    r = requests.post("https://api.dropboxapi.com/oauth2/token", data={
        "grant_type": "refresh_token",
        "refresh_token": DROPBOX_REFRESH,
        "client_id": DROPBOX_APP_KEY,
        "client_secret": DROPBOX_APP_SECRET,
    })
    return r.json()["access_token"]

def upload_to_dropbox(token, local_path, dropbox_path):
    import requests
    with open(local_path, "rb") as f:
        data = f.read()
    r = requests.post(
        "https://content.dropboxapi.com/2/files/upload",
        headers={
            "Authorization": f"Bearer {token}",
            "Dropbox-API-Arg": json.dumps({"path": dropbox_path, "mode": "overwrite"}),
            "Content-Type": "application/octet-stream",
        },
        data=data
    )
    if r.status_code != 200:
        raise Exception(f"Dropbox upload error: {r.status_code} {r.text[:200]}")
    
    # Создаём shared link
    r2 = requests.post(
        "https://api.dropboxapi.com/2/sharing/create_shared_link_with_settings",
        headers={"Authorization": f"Bearer {token}", "Content-Type": "application/json"},
        json={"path": dropbox_path, "settings": {}}
    )
    if r2.status_code == 200:
        url = r2.json().get("url", "")
        return url.replace("?dl=1", "?dl=0") if url else dropbox_path
    # Если ссылка уже есть
    if r2.status_code == 409:
        data2 = r2.json()
        existing = data2.get("error", {}).get("shared_link_already_exists", {}).get("metadata", {}).get("url", "")
        if existing:
            return existing.replace("?dl=1", "?dl=0")
    return dropbox_path

# ─── DOCX ─────────────────────────────────────────────────────────────────────
def split_long_paragraph(text, max_sentences=3, max_words=70):
    """Split a paragraph that is too long into shorter ones (max 3 sentences / 70 words)."""
    import re as _re
    sentences = _re.split(r'(?<=[.!?])\s+', text.strip())
    if len(sentences) <= max_sentences and len(text.split()) <= max_words:
        return [text]
    
    chunks = []
    current = []
    current_words = 0
    for sent in sentences:
        sent_words = len(sent.split())
        if current and (len(current) >= max_sentences or current_words + sent_words > max_words):
            chunks.append(' '.join(current))
            current = [sent]
            current_words = sent_words
        else:
            current.append(sent)
            current_words += sent_words
    if current:
        chunks.append(' '.join(current))
    return chunks if chunks else [text]


def text_to_html_txt(text, output_path):
    """
    Convert text to .txt with clean HTML tags for website posting.
    Format:
      <!-- Meta Title: ... -->
      <!-- Meta Description: ... -->
      <h1>...</h1>
      <p>...</p>
      <h2>...</h2>
      <table><thead><tr><th>...</th></tr></thead><tbody><tr><td>...</td></tr></tbody></table>
      <ul><li>...</li></ul>
      <ol><li>...</li></ol>
    """
    import re as _re

    def close_open(ctx, tag):
        if ctx['in_table']:
            ctx['rows'] and _flush_table(ctx)
        if ctx['in_ul']:
            ctx['out'].append('</ul>'); ctx['in_ul'] = False
        if ctx['in_ol']:
            ctx['out'].append('</ol>'); ctx['in_ol'] = False

    def _flush_table(ctx):
        rows = ctx['rows']
        if not rows:
            return
        out = ctx['out']
        out.append('<table>')
        # First row = header
        out.append('<thead>')
        out.append('<tr>' + ''.join(f'<th>{c}</th>' for c in rows[0]) + '</tr>')
        out.append('</thead>')
        if len(rows) > 1:
            out.append('<tbody>')
            for r in rows[1:]:
                out.append('<tr>' + ''.join(f'<td>{c}</td>' for c in r) + '</tr>')
            out.append('</tbody>')
        out.append('</table>')
        ctx['rows'] = []
        ctx['in_table'] = False

    def inline(s):
        """Convert inline markdown to HTML."""
        s = _re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', s)
        s = _re.sub(r'\*(.+?)\*', r'<em>\1</em>', s)
        s = _re.sub(r'`(.+?)`', r'<code>\1</code>', s)
        return s

    ctx = {'out': [], 'in_ul': False, 'in_ol': False, 'in_table': False, 'rows': []}
    lines = text.split('\n')

    for line in lines:
        s = line.strip()

        # Empty line
        if not s:
            if ctx['in_table']:
                _flush_table(ctx)
            if ctx['in_ul']:
                ctx['out'].append('</ul>'); ctx['in_ul'] = False
            if ctx['in_ol']:
                ctx['out'].append('</ol>'); ctx['in_ol'] = False
            continue

        # Meta comments
        if s.startswith('Meta Title:'):
            if ctx['in_table']: _flush_table(ctx)
            ctx['out'].append(f'<!-- Meta Title: {s[11:].strip()} -->')
            continue
        if s.startswith('Meta Description:'):
            if ctx['in_table']: _flush_table(ctx)
            ctx['out'].append(f'<!-- Meta Description: {s[17:].strip()} -->')
            continue

        # Headings H1/H2/H3/H4
        hm = _re.match(r'^(H[1-4]):\s*(.*)', s)
        if hm:
            if ctx['in_table']: _flush_table(ctx)
            if ctx['in_ul']: ctx['out'].append('</ul>'); ctx['in_ul'] = False
            if ctx['in_ol']: ctx['out'].append('</ol>'); ctx['in_ol'] = False
            level = hm.group(1).lower()
            ctx['out'].append(f'<{level}>{inline(hm.group(2))}</{level}>')
            continue

        # Markdown headings ## / ###
        mdh = _re.match(r'^(#{1,4})\s+(.*)', s)
        if mdh:
            if ctx['in_table']: _flush_table(ctx)
            if ctx['in_ul']: ctx['out'].append('</ul>'); ctx['in_ul'] = False
            if ctx['in_ol']: ctx['out'].append('</ol>'); ctx['in_ol'] = False
            level = min(len(mdh.group(1)), 4)
            ctx['out'].append(f'<h{level}>{inline(mdh.group(2))}</h{level}>')
            continue

        # Pipe table row
        if s.startswith('|') and '|' in s[1:]:
            if ctx['in_ul']: ctx['out'].append('</ul>'); ctx['in_ul'] = False
            if ctx['in_ol']: ctx['out'].append('</ol>'); ctx['in_ol'] = False
            cells = [c.strip() for c in s.strip('|').split('|')]
            # Skip separator |---|---|
            if all(_re.match(r'^[-: ]+$', c) for c in cells if c):
                continue
            ctx['in_table'] = True
            ctx['rows'].append(cells)
            continue
        else:
            if ctx['in_table']:
                _flush_table(ctx)

        # Unordered list
        if _re.match(r'^[-*\u2022]\s+', s):
            if ctx['in_ol']: ctx['out'].append('</ol>'); ctx['in_ol'] = False
            if not ctx['in_ul']:
                ctx['out'].append('<ul>'); ctx['in_ul'] = True
            item = _re.sub(r'^[-*\u2022]\s+', '', s)
            ctx['out'].append(f'<li>{inline(item)}</li>')
            continue

        # Ordered list
        nm = _re.match(r'^\d+[.):]\s+(.*)', s)
        if nm:
            if ctx['in_ul']: ctx['out'].append('</ul>'); ctx['in_ul'] = False
            if not ctx['in_ol']:
                ctx['out'].append('<ol>'); ctx['in_ol'] = True
            ctx['out'].append(f'<li>{inline(nm.group(1))}</li>')
            continue

        # Close lists
        if ctx['in_ul']: ctx['out'].append('</ul>'); ctx['in_ul'] = False
        if ctx['in_ol']: ctx['out'].append('</ol>'); ctx['in_ol'] = False

        # Skip lines that are just dividers or artifacts
        if _re.match(r'^[-=*_]{3,}$', s):
            continue
        # Skip prompt artifacts
        if s.startswith('===') or s.startswith('---') and len(s) < 6:
            continue

        # Regular paragraph
        ctx['out'].append(f'<p>{inline(s)}</p>')

    # Close any open blocks
    if ctx['in_table']: _flush_table(ctx)
    if ctx['in_ul']: ctx['out'].append('</ul>')
    if ctx['in_ol']: ctx['out'].append('</ol>')

    # Join with single newlines, remove consecutive blank lines
    result = '\n'.join(ctx['out'])
    result = _re.sub(r'\n{3,}', '\n\n', result)

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(result)


def text_to_docx(text, output_path):
    """Конвертируем текст в DOCX с правильными таблицами"""
    # Extract Meta Title and Meta Description from top of text
    import re as _re
    meta_title = None
    meta_desc = None
    meta_title_match = _re.search(r'(?m)^Meta Title:\s*(.+)$', text)
    if meta_title_match:
        meta_title = meta_title_match.group(1).strip()
        text = text[:meta_title_match.start()] + text[meta_title_match.end():]
    meta_desc_match = _re.search(r'(?m)^Meta Description:\s*(.+)$', text)
    if meta_desc_match:
        meta_desc = meta_desc_match.group(1).strip()
        text = text[:meta_desc_match.start()] + text[meta_desc_match.end():]
    # Aggressively strip Meta Title/Desc that may not have been caught
    text = _re.sub(r'(?m)^Meta Title:.*$\n?', '', text)
    text = _re.sub(r'(?m)^Meta Description:.*$\n?', '', text)
    # Strip ALL prompt markers Gemini may copy verbatim
    text = _re.sub(r'={3,}[^=\n]+={3,}\n?', '', text)              # ===TEXT===
    text = _re.sub(r'[═─━]{3,}[^\n]*[═─━]{3,}\n?', '', text)      # ═══TEXT═══
    text = _re.sub(r'(?m)^\s*[═=─━]{3,}\s*$\n?', '', text)       # standalone dividers
    text = _re.sub(r'(?m)^\s*-{5,}\s*$\n?', '', text)             # ----- dividers
    text = _re.sub(r'(?m)^\s*[*]{3,}\s*$\n?', '', text)           # *** dividers
    # Strip lines that are pure instruction text from prompt
    text = _re.sub(r'(?m)^--- Now write.*$\n?', '', text)
    text = _re.sub(r'(?m)^=== PAGE STRUCTURE.*$\n?', '', text)
    text = _re.sub(r'(?m)^IMPORTANT: Do NOT repeat.*$\n?', '', text)
    text = _re.sub(r'(?m)^MANDATORY ARTICLE STRUCTURE.*$\n?', '', text)
    text = _re.sub(r'(?m)^FORBIDDEN PATTERNS:.*$\n?', '', text)
    text = _re.sub(r'(?m)^DEEP REVIEW SECTIONS.*$\n?', '', text)
    # Clean up multiple blank lines
    text = _re.sub(r'\n{3,}', '\n\n', text)
    text = text.lstrip('\n')

    # Clean up markdown that Gemini sometimes outputs despite instructions
    # Remove **bold** markers — convert to plain text (Word styles handle formatting)
    text = _re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    # Remove *italic* markers
    text = _re.sub(r'\*(.+?)\*', r'\1', text)
    # Remove markdown # headings (should already be H1:/H2: format)
    text = _re.sub(r'^#{1,6}\s+', '', text, flags=_re.MULTILINE)
    # Strip FAQ letter/number prefixes: Q:, A:, Q1., Q1:, A1., 1), a), b), etc.
    text = _re.sub(r'(?m)^\s*[QqAa]\s*[:.]\s*', '', text)          # Q: / A: / q. / a.
    text = _re.sub(r'(?m)^\s*[QqAa]\d+[:.)]\s*', '', text)         # Q1: / A1. / Q2)
    text = _re.sub(r'(?m)^\s*\d+[).:]\s+(?=[A-ZÄÖÜA-Za-zА-Яа-я])', '', text)  # 1. / 1) before capital
    text = _re.sub(r'(?m)^\s*[a-z][).:]\s+', '', text)              # a) / b. prefix
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import re

        doc = Document()
        
        # Настройка шрифта
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        # Add Meta Title and Description at top of document if present
        if meta_title:
            p_mt = doc.add_paragraph()
            run = p_mt.add_run("Meta Title: ")
            run.bold = True
            run.font.size = Pt(10)
            run2 = p_mt.add_run(meta_title)
            run2.font.size = Pt(10)
            p_mt.paragraph_format.space_after = Pt(2)
        if meta_desc:
            p_md = doc.add_paragraph()
            run = p_md.add_run("Meta Description: ")
            run.bold = True
            run.font.size = Pt(10)
            run2 = p_md.add_run(meta_desc)
            run2.font.size = Pt(10)
            p_md.paragraph_format.space_after = Pt(8)
        if meta_title or meta_desc:
            doc.add_paragraph()  # separator line

        def add_table_border(table):
            """Добавляем границы таблице"""
            for row in table.rows:
                for cell in row.cells:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcBorders = OxmlElement("w:tcBorders")
                    for border_name in ["top","left","bottom","right","insideH","insideV"]:
                        border = OxmlElement(f"w:{border_name}")
                        border.set(qn("w:val"), "single")
                        border.set(qn("w:sz"), "4")
                        border.set(qn("w:color"), "000000")
                        tcBorders.append(border)
                    tcPr.append(tcBorders)

        lines = text.splitlines()
        i = 0
        while i < len(lines):
            line = lines[i]
            s = line.strip()

            # Пустая строка
            if not s:
                doc.add_paragraph()
                i += 1
                continue
            
            # Убираем ** вокруг строки (Gemini иногда оборачивает заголовки)
            if s.startswith("**") and s.endswith("**") and len(s) > 4:
                s = s[2:-2].strip()
                line = s

            # Заголовки H1/H2/H3
            if s.startswith("H1:"):
                doc.add_heading(s[3:].strip(), level=1)
                i += 1
                continue
            elif s.startswith("H2:") or s.startswith("H2 "):
                doc.add_heading(s[3:].strip(), level=2)
                i += 1
                continue
            elif s.startswith("H3:") or s.startswith("H3 "):
                doc.add_heading(s[3:].strip(), level=3)
                i += 1
                continue

            # Таблица в pipe-формате (с | или без ведущего |)
            is_pipe_row = (s.startswith("|") and "|" in s[1:]) or                           ("|" in s and s.count("|") >= 2 and not s.startswith("#") and not s.startswith("H"))
            if is_pipe_row:
                # Собираем все строки таблицы
                table_lines = []
                while i < len(lines):
                    row_line = lines[i].strip()
                    if not row_line:
                        break
                    has_pipe = (row_line.startswith("|") and "|" in row_line[1:]) or                                ("|" in row_line and row_line.count("|") >= 2)
                    if not has_pipe:
                        break
                    # Пропускаем separator строки (---|---|--- или |:---|:---|)
                    if re.match(r"^[\|\s]*[-:\s|]+[\|\s]*$", row_line) and "-" in row_line:
                        i += 1
                        continue
                    # Parse cells
                    if row_line.startswith("|"):
                        cells = [c.strip() for c in row_line.split("|")[1:-1]]
                    else:
                        cells = [c.strip() for c in row_line.split("|")]
                    if cells and any(c for c in cells):
                        table_lines.append(cells)
                    i += 1
                
                if table_lines:
                    max_cols = max(len(r) for r in table_lines)
                    table = doc.add_table(rows=len(table_lines), cols=max_cols)
                    table.style = "Table Grid"
                    
                    for ri, row_data in enumerate(table_lines):
                        for ci, cell_text in enumerate(row_data):
                            if ci < max_cols:
                                cell = table.cell(ri, ci)
                                cell.text = cell_text
                                if ri == 0:
                                    # Заголовок жирным
                                    for para in cell.paragraphs:
                                        for run in para.runs:
                                            run.bold = True
                    
                    add_table_border(table)
                    doc.add_paragraph()
                continue

            # Список
            if re.match(r"^[-*•]\s", s):
                item_text = s[2:].strip()
                # Long descriptive lines (> 120 chars) → split into Normal para, not bullet
                # Short clean items → List Bullet
                if len(item_text) > 120:
                    p = doc.add_paragraph(item_text, style="Normal")
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after = Pt(4)
                else:
                    doc.add_paragraph(item_text, style="List Bullet")
                i += 1
                continue

            # Нумерованный список
            if re.match(r"^\d+[.)]\s", s):
                text_content = re.sub(r"^\d+[.)]\s*", "", s)
                doc.add_paragraph(text_content, style="List Number")
                i += 1
                continue

            # Обычный параграф — auto-split if too long
            chunks = split_long_paragraph(s, max_sentences=3, max_words=70)
            for chunk in chunks:
                p = doc.add_paragraph(chunk)
                p.paragraph_format.space_after = Pt(6)
            i += 1

        doc.save(output_path)
        return True
    except ImportError as e:
        print(f"  docx import error: {e}")
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(text)
        return True
    except Exception as e:
        print(f"  docx error: {e}")
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(text)
        return True


# ─── Gemini ───────────────────────────────────────────────────────────────────

def check_readability(text):
    """
    Check sentence length distribution.
    Returns dict with stats and warnings.
    """
    import re
    sentences = re.split(r'(?<=[.!?])\s+', text)
    sentences = [s.strip() for s in sentences if len(s.strip().split()) >= 3]
    
    if not sentences:
        return {"ok": True, "sentences": 0}
    
    lengths = [len(s.split()) for s in sentences]
    avg = sum(lengths) / len(lengths)
    too_short = sum(1 for l in lengths if l <= 7)
    too_long = sum(1 for l in lengths if l >= 22)
    optimal = sum(1 for l in lengths if 8 <= l <= 21)
    
    total = len(lengths)
    short_pct = too_short / total * 100
    long_pct = too_long / total * 100
    optimal_pct = optimal / total * 100
    
    warnings = []
    if short_pct > 20:
        warnings.append(f"⚠️ {short_pct:.0f}% sentences too short (≤7 words) — text feels choppy")
    if long_pct > 15:
        warnings.append(f"⚠️ {long_pct:.0f}% sentences too long (≥22 words) — hurt readability")
    if optimal_pct < 60:
        warnings.append(f"⚠️ Only {optimal_pct:.0f}% sentences in optimal range (8-21 words)")
    
    return {
        "ok": len(warnings) == 0,
        "sentences": total,
        "avg_length": round(avg, 1),
        "optimal_pct": round(optimal_pct, 1),
        "too_short_pct": round(short_pct, 1),
        "too_long_pct": round(long_pct, 1),
        "warnings": warnings,
    }


def calculate_lix(text):
    """
    LIX (Läsbarhetsindex) — universal readability score for non-English.
    Works for: IT, PT, DE, NL, PL, ES, FR etc.
    Formula: LIX = A/B + C*100/A
    where A=words, B=sentences, C=words with >6 letters
    
    Score interpretation:
    <25 = very easy (children's books)
    25-35 = easy
    35-45 = medium (newspapers)
    45-55 = difficult
    >55 = very difficult (academic) — FLAG as too complex
    """
    import re
    if not text or len(text) < 50:
        return None
    
    words = text.split()
    A = len(words)
    if A == 0:
        return None
    
    sentences = re.split(r'[.!?]+', text)
    sentences = [s.strip() for s in sentences if len(s.strip()) > 3]
    B = max(len(sentences), 1)
    
    C = sum(1 for w in words if len(re.sub(r'[^a-zA-ZÀ-ÿа-яёА-ЯЁ]', '', w)) > 6)
    
    lix = (A / B) + (C * 100 / A)
    
    ok = lix <= 55
    if lix > 55:
        warning = f"⚠️ LIX={lix:.1f} — text too complex (>55), simplify sentences and vocabulary"
    elif lix > 45:
        warning = f"ℹ️ LIX={lix:.1f} — difficult (target <45 for casino content)"
    else:
        warning = ""
    
    return {
        "score": round(lix, 1),
        "ok": ok,
        "words": A,
        "sentences": B,
        "long_words": C,
        "warning": warning,
        "label": "very easy" if lix < 25 else "easy" if lix < 35 else "medium" if lix < 45 else "difficult" if lix < 55 else "VERY DIFFICULT",
    }


def calculate_fkgl(text):
    """
    Flesch-Kincaid Grade Level — for English texts.
    FKGL = 0.39 × (words/sentences) + 11.8 × (syllables/words) - 15.59
    Target: Grade 7–9 (broadly readable)
    """
    import re
    if not text or len(text) < 50:
        return None
    
    words = text.split()
    A = len(words)
    if A == 0:
        return None
    
    sentences = re.split(r'[.!?]+', text)
    sentences = [s.strip() for s in sentences if len(s.strip()) > 3]
    B = max(len(sentences), 1)
    
    def count_syllables(word):
        word = word.lower().strip(".,!?;:\"'")
        if len(word) <= 3:
            return 1
        word = re.sub(r'[^a-z]', '', word)
        count = len(re.findall(r'[aeiouy]+', word))
        if word.endswith('e') and count > 1:
            count -= 1
        return max(count, 1)
    
    total_syllables = sum(count_syllables(w) for w in words)
    
    fkgl = 0.39 * (A / B) + 11.8 * (total_syllables / A) - 15.59
    fkgl = max(0, fkgl)
    
    ok = 7 <= fkgl <= 9
    if fkgl < 7:
        note = f"ℹ️ FKGL={fkgl:.1f} — very easy (Grade {fkgl:.0f})"
    elif fkgl <= 9:
        note = f"✅ FKGL={fkgl:.1f} — target range Grade 7-9"
    else:
        note = f"⚠️ FKGL={fkgl:.1f} — too complex (Grade {fkgl:.0f}, target 7-9)"
    
    return {
        "score": round(fkgl, 1),
        "ok": ok,
        "words": A,
        "sentences": B,
        "syllables": total_syllables,
        "warning": note,
        "label": f"Grade {fkgl:.0f}",
    }


def check_keyword_density(text, keyword, word_count=None):
    """
    Check if keyword appears with correct density: 1-3 per 100 words.
    Returns density stats and warnings.
    """
    if not keyword or not text:
        return {}
    import re
    total_words = len(text.split()) if not word_count else word_count
    # Case-insensitive count
    count = len(re.findall(re.escape(keyword.lower()), text.lower()))
    density_per_100 = (count / total_words * 100) if total_words else 0
    
    ok = 1.0 <= density_per_100 <= 3.0
    warning = ""
    if density_per_100 < 1.0:
        warning = f"⚠️ Keyword '{keyword}' underused: {count}x / {density_per_100:.2f}% (target: 1-3%)"
    elif density_per_100 > 3.5:
        warning = f"⚠️ Keyword '{keyword}' overused: {count}x / {density_per_100:.2f}% (target: 1-3%) — may trigger stuffing filter"
    
    return {
        "keyword": keyword,
        "count": count,
        "density_pct": round(density_per_100, 2),
        "ok": ok,
        "warning": warning,
    }


def check_paragraph_length(text):
    """
    Analyze paragraph sentence counts.
    Target: 2-4 sentences per paragraph (working hypothesis — calibrate vs competitors).
    """
    import re
    # Split into paragraphs by blank lines or block endings
    paragraphs = [p.strip() for p in re.split(r'\n{2,}', text) if p.strip()]
    # Filter out headings and very short lines
    content_paras = [p for p in paragraphs if not p.startswith('H') and len(p.split()) > 10]
    
    if not content_paras:
        return {}
    
    sent_counts = []
    for para in content_paras:
        sents = len(re.split(r'(?<=[.!?])\s+', para))
        sent_counts.append(sents)
    
    avg = sum(sent_counts) / len(sent_counts)
    too_long = sum(1 for s in sent_counts if s > 4)
    too_short = sum(1 for s in sent_counts if s < 2)
    optimal = sum(1 for s in sent_counts if 2 <= s <= 4)
    
    warnings = []
    if too_long / len(sent_counts) > 0.25:
        warnings.append(f"⚠️ {too_long}/{len(sent_counts)} paragraphs > 4 sentences — may hurt readability")
    
    return {
        "total_paras": len(content_paras),
        "avg_sentences": round(avg, 1),
        "optimal_pct": round(optimal / len(sent_counts) * 100, 1),
        "too_long_count": too_long,
        "distribution": sorted(set(sent_counts)),
        "warnings": warnings,
    }

def _detect_wrong_language(text: str, target: str) -> str:
    """
    Detect if text contains significant amount of wrong language.
    Returns reason string if wrong language detected, empty string if OK.
    Uses character-based heuristic + common-word markers.
    """
    import re as _rel
    target = (target or "EN").upper()[:2]

    # Strip code blocks, URLs, brand names (uppercase words), numbers
    clean = _rel.sub(r'https?://\S+', '', text)
    clean = _rel.sub(r'\b[A-Z][a-z]*[A-Z]\w*\b', '', clean)  # CamelCase brands
    clean = _rel.sub(r'\b\d+\b', '', clean)
    clean = _rel.sub(r'[^\w\s]', ' ', clean)

    # Common words per language (high-frequency markers + technical EN terms often left in translations)
    markers = {
        'EN': {'the','and','for','with','that','this','from','have','will','you','our','are','was','what','how',
               'mobile','experience','review','rating','bonus','bonuses','welcome','payment','payments',
               'methods','security','licensing','promotions','players','best','top','guide','ranking',
               'responsible','gambling','browser','deposit','withdrawal','games','support','safe','fast',
               'about','why','choose','get','started','login','account','platform','odds','selection',
               'available','products','sports','online','new','good','great','features','vs','sign','up'},
        'NL': {'de','het','een','van','en','is','dat','in','op','voor','met','zijn','niet','aan','ook'},
        'IT': {'il','la','le','di','che','un','una','per','con','del','della','sono','non','si','come'},
        'DE': {'der','die','das','und','ist','nicht','von','zu','mit','auf','für','eine','einen','sind','als'},
        'PL': {'jest','nie','sie','na','do','w','z','i','to','co','oraz','tak','dla','te','jak','tylko'},
        'FR': {'le','la','les','de','et','est','un','une','pour','avec','que','dans','sur','pas','ce'},
        'ES': {'el','la','los','las','de','y','es','un','una','para','con','que','en','por','su'},
        'PT': {'o','a','os','as','de','e','é','um','uma','para','com','que','em','por','no'},
        'HU': {'a','az','is','hogy','egy','nem','van','ezt','csak','meg','és','de','vagy','így','ha'},
        'SV': {'och','att','en','som','för','med','av','på','är','det','inte','till','den','har','vi'},
        'RU': {'и','в','не','что','как','это','на','с','по','для','от','из','при','но','то'},
        'UK': {'і','в','не','що','як','це','на','з','по','для','від','із','при','але','то'},
    }

    words = [w.lower() for w in clean.split() if len(w) >= 2]
    if len(words) < 100:
        return ""  # too short to judge

    # Count marker hits per language
    counts = {}
    for lang_code, mset in markers.items():
        counts[lang_code] = sum(1 for w in words if w in mset)

    target_count = counts.get(target, 0)
    other_langs = {k: v for k, v in counts.items() if k != target and v > 0}

    # For non-EN targets, recalculate EN count using ONLY strong-EN markers (not technical words)
    # This avoids false positives from words like 'casino', 'bonus', 'mobile' which are international
    if target != 'EN':
        strong_en_only = {'the','and','for','with','that','this','from','have','will','you','our',
                         'are','was','what','how','can','any','one','two','also','just','more',
                         'than','about','into','only','very','most','best','your','their','they'}
        en_strong_count = sum(1 for w in words if w in strong_en_only)
        if 'EN' in other_langs:
            other_langs['EN'] = en_strong_count
            counts['EN'] = en_strong_count

    # If another language has dominant marker count — BIG PROBLEM
    for lang_code, cnt in other_langs.items():
        if cnt > target_count * 1.2 and cnt > 10:
            return f"detected {lang_code} ({cnt} markers) dominates over {target} ({target_count})"
        # If another language has >50% of target's count and 20+ markers — mix
        if target_count > 0 and cnt / max(target_count, 1) > 0.5 and cnt > 20:
            return f"{lang_code} markers ({cnt}) too high vs {target} ({target_count})"

    # If target language has very few markers — probably wrong language entirely
    if target_count < 5 and len(words) > 300:
        top_other = sorted(other_langs.items(), key=lambda x: -x[1])[:2]
        if top_other and top_other[0][1] > 10:
            return f"very few {target} markers ({target_count}), dominant: {top_other[0][0]} ({top_other[0][1]})"

    # Paragraph-level check: detect blocks of non-target-language text
    # Split text into paragraphs, check each for dominant language
    paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
    foreign_paragraphs = 0
    for para in paragraphs:
        pwords = [w.lower() for w in _rel.sub(r'[^\w\s]', ' ', para).split() if len(w) >= 2]
        if len(pwords) < 20:
            continue
        p_target = sum(1 for w in pwords if w in markers.get(target, set()))
        p_other_max = 0
        p_other_lang = None
        for lc, ms in markers.items():
            if lc == target:
                continue
            c = sum(1 for w in pwords if w in ms)
            if c > p_other_max:
                p_other_max = c
                p_other_lang = lc
        # Paragraph-level: if other lang clearly dominates (strict to avoid false positives)
        if p_other_max > p_target * 3 and p_other_max >= 6:
            foreign_paragraphs += 1

    if foreign_paragraphs >= 2:
        return f"found {foreign_paragraphs} paragraph(s) in wrong language (target={target})"

    # HEADING-level check: find ALL H1/H2/H3/H4 and ensure each is in target language
    # Match both 'H2: text' format and '<h2>text</h2>' HTML format
    heading_patterns = [
        _rel.compile(r'(?m)^H[1-4]:\s*(.+)$'),
        _rel.compile(r'<h[1-4][^>]*>([^<]+)</h[1-4]>', _rel.IGNORECASE),
    ]
    # International words — acceptable in any language (don't trigger language mismatch)
    INTERNATIONAL = {
        'top','casino','casinos','bonus','online','promo','jackpot','live','slot','slots',
        'mobile','app','vip','crypto','bitcoin','poker','blackjack','roulette','baccarat',
        'rtp','review','test','play','start','win','game','games','sport','esports',
        'paypal','visa','mastercard','skrill','neteller','ideal','blik','klarna','revolut',
        'aams','adm','mga','ksa','ukgc','curacao','malta','gibraltar',
        'free','spin','spins','reload','cashback','wager','wagering',
        'classifica','analisi','migliori','recensione','recensioni',  # IT specific
        'najlepsze','recenzje',  # PL specific
        'beste','recensies',  # NL specific
        'beste','bewertung','bewertungen',  # DE specific
    }

    foreign_headings = []
    for pat in heading_patterns:
        for m in pat.finditer(text):
            h_text = m.group(1).strip()
            # Strip inline HTML, brands (CamelCase), numbers
            h_clean = _rel.sub(r'<[^>]+>', '', h_text)
            h_clean = _rel.sub(r'\b[A-Z][a-z]*[A-Z]\w*\b', '', h_clean)  # CamelCase brands
            h_clean = _rel.sub(r'\b\d+\b', '', h_clean)
            # Skip Title Case multi-word capitalized names (likely brand names)
            h_words_raw = h_clean.split()
            h_words = [w.lower() for w in _rel.sub(r'[^\w\s]', ' ', h_clean).split() if len(w) >= 2]
            # Filter out international words — they don't count as foreign
            h_words = [w for w in h_words if w not in INTERNATIONAL]
            if len(h_words) < 2:
                continue
            h_target = sum(1 for w in h_words if w in markers.get(target, set()))
            # Special check: if target is NOT EN, look for STRONG EN-only marker words
            if target != 'EN':
                # Count only strong EN words (the/and/for/with/that/this etc, not technical)
                strong_en_only = {'the','and','for','with','that','this','from','have','will','you','our',
                                  'are','was','what','how','can','any','one','two','also','just','more',
                                  'than','about','into','only','very','most','best','play','your'}
                en_in_heading = sum(1 for w in h_words if w in strong_en_only)
                if en_in_heading >= 2 and en_in_heading > h_target:
                    foreign_headings.append((h_text[:80], 'EN', en_in_heading, h_target))
                    continue
            h_other_max = 0
            h_other_lang = None
            for lc, ms in markers.items():
                if lc == target:
                    continue
                # Skip 'EN' here — already handled above with stricter check
                if lc == 'EN':
                    continue
                c = sum(1 for w in h_words if w in ms)
                if c > h_other_max:
                    h_other_max = c
                    h_other_lang = lc
            # Heading has notably more foreign markers than target
            if h_other_max >= 2 and h_other_max > h_target:
                foreign_headings.append((h_text[:80], h_other_lang, h_other_max, h_target))

    if foreign_headings:
        examples = '; '.join(f'"{h[0]}" ({h[1]}={h[2]}, {target}={h[3]})' for h in foreign_headings[:3])
        return f"{len(foreign_headings)} heading(s) in wrong language: {examples}"

    return ""


async def generate_text(client, prompt, idx, task=None, serp_targets=None):
    """Один запрос к Gemini"""
    print(f"  [{idx+1}] Генерирую...", flush=True)
    payload = {
        "model": _get_generation_model(),
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.7,
        "max_tokens": 32768,
    }
    for attempt in range(3):
        try:
            r = await client.post(
                GEMINI_URL,
                headers={"Authorization": f"Bearer {get_next_or_key()}", "Content-Type": "application/json"},
                json=payload,
                timeout=120.0
            )
            if r.status_code == 503:
                print(f"  [{idx+1}] 503, retry {attempt+1}/3...", flush=True)
                await asyncio.sleep(5 * (attempt + 1))
                continue
            r.raise_for_status()
            data = r.json()
            # OpenRouter response format
            choice = data["choices"][0]
            text = choice["message"]["content"]
            finish_reason = choice.get("finish_reason", "stop")
            if finish_reason == "MAX_TOKENS":
                print(f"  [{idx+1}] ⚠️  Text truncated (MAX_TOKENS) — attempting continuation...", flush=True)
                # Continue generation from where it stopped
                continuation_payload = {
                    "contents": [
                        {"role": "user", "parts": [{"text": prompt}]},
                        {"role": "model", "parts": [{"text": text}]},
                        {"role": "user", "parts": [{"text": "Continue writing from exactly where you stopped. Do not repeat any content. Continue seamlessly:"}]},
                    ],
                    "generationConfig": {"temperature": 0.7, "maxOutputTokens": 65536}
                }
                try:
                    r2 = await client.post(
                        f"{GEMINI_URL}?key={GEMINI_KEY}",
                        json=continuation_payload, timeout=120.0
                    )
                    if r2.status_code == 200:
                        data2 = r2.json()
                        continuation = data2["candidates"][0]["content"]["parts"][0]["text"]
                        text = text + "\n" + continuation
                        print(f"  [{idx+1}] ✅ Continuation added (+{len(continuation.split())} words)", flush=True)
                except Exception as cont_e:
                    print(f"  [{idx+1}] ⚠️  Continuation failed: {cont_e}", flush=True)

            word_count_actual = len(text.split())
            readability = check_readability(text)
            # Check primary keyword density
            primary_kw = (task or {}).get("keywords", [""])[0] if (task or {}).get("keywords") else ""
            kw_check = check_keyword_density(text, primary_kw, word_count_actual) if primary_kw else {}
            # Readability scores: LIX (universal) and FKGL (English)
            lang = (task or {}).get("language", "en") if task else "en"
            lix_result = calculate_lix(text)
            fkgl_result = calculate_fkgl(text) if lang.startswith("en") else None
            readability_info = ""
            if lix_result:
                lix_flag = " ⚠️COMPLEX" if not lix_result["ok"] else ""
                readability_info += f" | LIX={lix_result['score']}{lix_flag}"
            if fkgl_result:
                fkgl_flag = " ⚠️" if not fkgl_result["ok"] else ""
                readability_info += f" | FKGL={fkgl_result['score']}{fkgl_flag}"
            kw_ok = "" if kw_check.get("ok") else "⚠️"
            kw_info = f" | kw: {kw_check.get('count',0)}x/{kw_check.get('density_pct',0):.1f}%{kw_ok}" if kw_check else ""
            print(f"  [{idx+1}] ✅ {word_count_actual} слів | avg sent: {readability['avg_length']}w | optimal: {readability['optimal_pct']}%{kw_info}{readability_info}", flush=True)
            para_stats = check_paragraph_length(text)
            if para_stats:
                print(f"  [{idx+1}] 📊 Para stats: avg={para_stats['avg_sentences']}s/para | optimal={para_stats['optimal_pct']}% | total={para_stats['total_paras']}", flush=True)
            
            all_warnings = (
                readability.get("warnings", []) +
                ([lix_result["warning"]] if lix_result and not lix_result["ok"] and lix_result["warning"] else []) +
                ([fkgl_result["warning"]] if fkgl_result and not fkgl_result["ok"] else []) + 
                ([kw_check.get("warning")] if kw_check.get("warning") else []) +
                para_stats.get("warnings", [])
            )
            for w in all_warnings:
                if w: print(f"  [{idx+1}] {w}", flush=True)

            # 🚨 H1 CHECK — warning only (not blocking)
            import re as _reH1
            has_h1 = (
                bool(_reH1.search(r'(?m)^H1:\s+\S', text)) or
                bool(_reH1.search(r'<h1[^>]*>\S', text, _reH1.IGNORECASE)) or
                bool(_reH1.search(r'(?m)^#\s+\S', text))
            )
            if not has_h1:
                print(f"  [{idx+1}] ⚠️ NO H1 FOUND — warning only (not blocking)", flush=True)

            # 🚨 LANGUAGE CHECK — REJECT if wrong language detected
            target_lang = (task or {}).get("language", "EN").upper()[:2]
            lang_issue = _detect_wrong_language(text, target_lang)
            if lang_issue:
                print(f"  [{idx+1}] ❌ WRONG LANGUAGE — {lang_issue} (target: {target_lang}), retry {attempt+1}/3", flush=True)
                if attempt < 2:
                    await asyncio.sleep(3)
                    prompt = (
                        f"⚠️ ПРЕДЫДУЩИЙ ТЕКСТ ОТКЛОНЕН — НАЙДЕНЫ ЧУЖИЕ ЯЗЫКИ!\n"
                        f"🚨 ПИШИ ВЕСЬ ТЕКСТ ТОЛЬКО НА {target_lang}. НИ ОДНОГО СЛОВА НА ДРУГИХ ЯЗЫКАХ.\n"
                        f"ОБЯЗАТЕЛЬНО: заголовки, таблицы, FAQ, CTA — все на {target_lang}.\n"
                        f"Исключения только: бренды (имена казино), URL, валюты.\n\n"
                        + prompt
                    )
                    continue
            # ── v3: Post-generation content audit & auto-fix ──────────────
            if serp_targets:
                keyword = (task or {}).get("keywords", [""])
                keyword = keyword[0] if keyword else ""
                audit = analyze_content(text, task or {}, serp_targets)
                print(f"  [{idx+1}] {audit['report']}", flush=True)
                fixable = [i for i in audit["issues"] if i["fixable"]]
                if fixable:
                    print(f"  [{idx+1}] 🔧 Auto-fixing {len(fixable)} critical issue(s)...", flush=True)
                    text = fix_content_issues(text, fixable, task or {}, keyword)
                    # Re-audit after fix
                    audit2 = analyze_content(text, task or {}, serp_targets)
                    still = [i for i in audit2["issues"] if i["fixable"]]
                    if not still:
                        print(f"  [{idx+1}] ✅ All fixable issues resolved", flush=True)
                    else:
                        for iss in still:
                            print(f"  [{idx+1}] ⚠️  Still: {iss['label']}={iss['value']} (target={_range_str(iss['target'])})", flush=True)
            return text
        except Exception as e:
            if attempt < 2:
                print(f"  [{idx+1}] Retry {attempt+1}/3: {e}", flush=True)
                await asyncio.sleep(3)
            else:
                print(f"  [{idx+1}] ❌ Ошибка: {e}", flush=True)
                return None

# ─── Промпты ──────────────────────────────────────────────────────────────────
# Structure variations for rotating H2 block order across texts
STRUCTURE_VARIANTS = [
    # Variant 0: Ranking first, then bonuses, then mobile, then payments
    ["ranking", "bonuses", "mobile", "payments", "responsible", "conclusion", "faq"],
    # Variant 1: Bonuses first, then ranking, then payments, then mobile
    ["bonuses", "ranking", "payments", "mobile", "conclusion", "responsible", "faq"],
    # Variant 2: Mobile first (mobile-first angle), then ranking, then bonuses
    ["mobile", "ranking", "bonuses", "payments", "responsible", "conclusion", "faq"],
    # Variant 3: Payments/trust first, then ranking, then bonuses
    ["payments", "responsible", "ranking", "bonuses", "mobile", "conclusion", "faq"],
    # Variant 4: FAQ first (Q&A format), then ranking, then bonuses
    ["faq_intro", "ranking", "bonuses", "mobile", "payments", "responsible", "conclusion"],
]

# Brand description style variations
BRAND_DESC_STYLES = [
    "analytical",   # Focus on data: RTP, payout speed, license details
    "user_story",   # First-person experience angle: what a player feels
    "comparative",  # Compare to competitors: better than X because...
    "feature_list", # Structured: key features, pros/cons table
    "narrative",    # Story-driven: brand history, reputation
]



async def _pregen_unique_titles(task: dict, count: int, serp_targets = None) -> list[dict]:
    """
    Предварительно генерирует уникальные Meta Title + H1 + Meta Description
    для всего батча одним запросом к AI.
    Возвращает список: [{"meta_title": "...", "h1": "...", "meta_desc": "..."}]
    """
    import httpx as _hx_t, json as _json_t

    keyword  = task.get("keywords", [""])[0] if task.get("keywords") else task.get("keyword", "")
    keywords = task.get("keywords", [])
    geo      = task.get("geo", "")
    lang     = task.get("language", "en")
    casinos  = task.get("casinos", [])
    task_type = task.get("article_type", "text_review")
    import datetime; year = datetime.datetime.now().year

    # Конкурентные titles из serp_targets
    comp_titles = []
    if serp_targets and isinstance(serp_targets, dict):
        deep = serp_targets.get("deep") or {}
        comp_titles = deep.get("sample_titles", [])[:5]
    comp_titles_str = "\n".join(f'  - "{t}"' for t in comp_titles) if comp_titles else "  (нет данных)"

    kw_str = ", ".join(keywords[:8]) if keywords else keyword
    casinos_str = ", ".join(casinos[:5]) if casinos else ""

    prompt = f"""Generate {count} UNIQUE and VARIED Meta Titles, H1 headings, and Meta Descriptions for SEO articles.

TASK:
- Topic: {keyword}
- Language: {lang} (ALL output must be in this language)
- GEO: {geo}
- Year: {year}
- Article type: {task_type}
- Keywords: {kw_str}
{"- Brands to feature: " + casinos_str if casinos_str else ""}

COMPETITOR TITLES (for reference/inspiration, DO NOT copy):
{comp_titles_str}

REQUIREMENTS:
1. Each Meta Title must be UNIQUE — no two titles can be similar
2. Each H1 must be UNIQUE — varied structure (keyword first / keyword middle / keyword end)
3. Include primary keyword "{keyword}" in every title and H1
4. Include GEO "{geo}" and year "{year}" in every title
5. Length: Meta Title 50-60 chars, H1 40-70 chars, Meta Description 130-155 chars
6. Language: ALL text must be in {lang} — NO English unless lang=en
7. Vary the phrasing: use different words, structures, angles

OUTPUT FORMAT (JSON array only, no markdown):
[
  {{"meta_title": "...", "h1": "...", "meta_desc": "..."}},
  {{"meta_title": "...", "h1": "...", "meta_desc": "..."}}
]
Generate exactly {count} objects."""

    api_key = task.get("api_key") or task.get("gemini_key") or ""

    # Используем тот же API что и основная генерация
    try:
        import httpx as _hx2
        import re as _re_t
        async with _hx2.AsyncClient() as _cl:
            resp = await generate_text(_cl, prompt, 0, task=task)
        if resp:
            json_match = _re_t.search(r'\[.*?\]', resp, _re_t.DOTALL)
            if json_match:
                titles = _json_t.loads(json_match.group(0))
                if isinstance(titles, list) and len(titles) >= count:
                    print(f"  ✅ Pre-generated {len(titles)} unique titles")
                    for j, t in enumerate(titles[:count]):
                        _mt = t.get("meta_title", "")[:60]
                        print("    [" + str(j+1) + "] " + _mt)
                    return titles[:count]
    except Exception as e:
        print(f"  ⚠️ Title pre-gen failed: {e} — AI will generate its own")

    return []  # Fallback: пустой список, промпт сам сгенерирует



def build_prompt(task, article_type, casino=None, text_index=0, serp_targets=None, pre_title=None):
    geo = task.get("geo", "UK")
    from datetime import datetime as _dt
    current_year = _dt.now().year  # Always use actual current year
    lang = task.get("language", "EN")
    keywords = task.get("keywords", [])
    kw_str = ", ".join(keywords[:5])
    word_count = int(task.get("word_count") or 3000)  # safe int conversion
    requirements = task.get("requirements", "")
    all_casinos = task.get("casinos", [])
    # Brand source rule:
    # - If casinos specified in task → use ONLY those brands
    # - If no casinos specified → extract brands from competitor pages
    brands_from_task = bool(all_casinos)
    if not all_casinos:
        # Try to extract brand names from competitor structure analysis
        struct_data = task.get("competitor_structure", {})
        competitor_brands = task.get("competitor_brands", [])
        if competitor_brands:
            all_casinos = competitor_brands[:10]
            print(f"  ℹ️ No brands in task — using {len(all_casinos)} brands from competitor analysis")
    casinos_str = ", ".join(all_casinos) if all_casinos else ""
    task_type = task.get("article_type", "text_review")

    # Calculate keyword density targets
    kw_list = task.get("keywords", [])
    primary_kw = kw_list[0] if kw_list else ""
    secondary_kws = kw_list[1:5] if len(kw_list) > 1 else []
    # 1-3 per 100 words = for word_count, primary key appears N times
    kw_min = max(3, int(word_count / 100 * 1))
    kw_max = max(6, int(word_count / 100 * 2.5))

    # Structure and brand style variation based on text_index
    struct_variant = STRUCTURE_VARIANTS[text_index % len(STRUCTURE_VARIANTS)]
    brand_style = BRAND_DESC_STYLES[text_index % len(BRAND_DESC_STYLES)]
    h1_position = ["START", "MIDDLE", "END", "START", "MIDDLE"][text_index % 5]

    # Build structure instruction
    # NOTE: labels below are INTENTS (in English for clarity), bot MUST translate H2 titles into {lang}
    struct_labels = {
        "ranking": f"[SECTION: RANKING] — write H2 in {lang}: Top {max(5, len(all_casinos) if all_casinos else 5)} casinos ranking with table + H3 per brand (ALL heading text in {lang}, NOT English)",
        "bonuses": f"[SECTION: BONUSES] — write H2 in {lang} about bonuses/promotions for {geo} players (bonus types, wagering, tips)",
        "mobile": f"[SECTION: MOBILE] — write H2 in {lang} about mobile casino experience (app vs browser, best mobile picks)",
        "payments": f"[SECTION: PAYMENTS] — write H2 in {lang} about payment methods in {geo} (include table: method/deposit time/withdrawal/fees/min)",
        "responsible": f"[SECTION: RESPONSIBLE GAMBLING] — write H2 in {lang} about responsible play (self-exclusion, limits, support)",
        "conclusion": f"[SECTION: CONCLUSION] — write H2 in {lang} for summary + final picks (200+ words)",
        "faq": f"[SECTION: FAQ] — write H2 in {lang} (ОБЯЗАТЕЛЬНО минимум 16 вопросов с развёрнутыми ответами 3-5 предложений каждый). Вопросы должны включать все ключевые слова. Questions and answers ALL in {lang}",
        "faq_intro": f"[SECTION: QUICK FAQ] — write H2 in {lang} (FAQ-first format, top 3 most common questions answered first)",
    }
    struct_order = "\n".join([f"  {j+1}. {struct_labels.get(s, s)}" for j, s in enumerate(struct_variant)])

    brand_style_instructions = {
        "analytical": "Describe each brand using ANALYTICAL style: focus on RTP data, payout speed, license details, game count. Use numbers and facts.",
        "user_story": "Describe each brand using USER STORY style: write from perspective of a player. What do they experience? How does registration feel? What stands out?",
        "comparative": "Describe each brand using COMPARATIVE style: explicitly compare each brand to at least one competitor. Use phrases like 'unlike X', 'better than average', 'stands out because'.",
        "feature_list": "Describe each brand using FEATURE LIST style: lead with a bullet list of 4-5 key features before the paragraph description. Show pros/cons explicitly.",
        "narrative": "Describe each brand using NARRATIVE style: tell a short story about the brand — its reputation, what makes it memorable, what type of player it's perfect for.",
    }
    brand_desc_rule = brand_style_instructions.get(brand_style, brand_style_instructions["analytical"])

    # Inject TF target from competitor analysis
    tf_data = task.get("competitor_tf")
    tf_rule = ""
    if tf_data:
        tf_rule = (
            f"\n▸ COMPETITOR TF ANALYSIS — match these metrics:\n"
            f"  - Competitors use '{tf_data['keyword']}' avg {tf_data['avg_count']:.1f}x in {tf_data['avg_words']:.0f} words\n"
            f"  - Average TF: {tf_data['avg_tf_per_100']:.2f}% (per 100 words)\n"
            f"  - YOUR TARGET: {tf_data['recommendation']}\n"
            f"  - TF range: {tf_data['target_tf_min']}%–{tf_data['target_tf_max']}% per 100 words\n"
        )
    
    # Inject competitor structure analysis
    struct_data = task.get("competitor_structure")
    struct_rule = ""
    if struct_data:
        h2_samples = " | ".join(f'"{h}"' for h in struct_data.get("top_h2_samples", [])[:4])
        topics = ", ".join(struct_data.get("common_topics", [])[:6])
        struct_rule = (
            f"\n▸ COMPETITOR STRUCTURE ANALYSIS ({struct_data['sites_analyzed']} sites) — mirror this structure:\n"
            f"  - Use {struct_data['avg_h2_count']} H2 sections (avg from top competitors)\n"
            f"  - Use {struct_data['avg_h3_count']} H3 subsections per article\n"
            f"  - {'MUST include comparison table — ' + str(struct_data['table_pct']) + '% of competitors use tables' if struct_data['table_pct'] >= 50 else 'Tables optional — only ' + str(struct_data['table_pct']) + '% use them'}\n"
            f"  - {'MUST use bullet lists — ' + str(struct_data['list_pct']) + '% of competitors use lists' if struct_data['list_pct'] >= 60 else 'Lists used by ' + str(struct_data['list_pct']) + '%'}\n"
            f"  - {'MUST include callout/highlight boxes — ' + str(struct_data['callout_pct']) + '% use them' if struct_data['callout_pct'] >= 40 else ''}\n"
            + (f"  - TOP H2 TOPICS used by competitors (inspire yours): {h2_samples}\n" if h2_samples else "")
            + (f"  - COMMON TOPICS to cover: {topics}\n" if topics else "")
            + (f"  - {'CRITICAL: competitors use first-person experience phrases — YOU MUST TOO: We tested, Our review found, During our analysis...' if struct_data['has_experience_phrases'] else ''}\n")
            + f"  - Mirror their content depth and block variety — not just length\n"
        )


    # Inject DEEP competitor analysis (LSI, links, paragraphs, FAQ, on-page SEO)
    deep_data = task.get("competitor_deep")
    deep_rule = ""
    if deep_data:
        top_lsi_str = ", ".join(f"{w}×{c}" for w, c in deep_data.get('top_lsi_terms', [])[:15])
        long_tail_str = "; ".join(f'"{p}"' for p in deep_data.get('long_tail_samples', [])[:8])
        common_dom_str = ", ".join(deep_data.get('common_external_domains', [])[:8])
        sample_titles = " | ".join(f'"{t}"' for t in deep_data.get('sample_titles', [])[:2])

        deep_rule = (
            f"\n═══ 🔬 DEEP COMPETITOR ANALYSIS (top {deep_data['sites_analyzed']} sites) ═══\n"
            f"⚠️ Use these as TARGETS to MATCH or BEAT — do not invent random numbers.\n\n"

            f"▸ CONTENT STRUCTURE (match these):\n"
            f"  - Body words: {deep_data['avg_body_words']} avg (you should have similar)\n"
            f"  - Paragraphs: {deep_data['avg_paragraph_words']}w avg, max {deep_data['max_paragraph_words']}w — keep yours ≤ 60w (better than competitors)\n"
            f"  - H2 count: {deep_data['avg_h2_count']} (target {int(deep_data['avg_h2_count'])}–{int(deep_data['avg_h2_count'])+3})\n"
            f"  - H3 count: {deep_data['avg_h3_count']} (target {int(deep_data['avg_h3_count'])}–{int(deep_data['avg_h3_count'])+5})\n"
            f"  - Tables: {deep_data['avg_table_count']} (avg {deep_data['avg_table_rows']} rows each) — INCLUDE at least {max(int(deep_data['avg_table_count']), 2)} tables\n"
            f"  - Lists: {deep_data['avg_list_count']} (avg {deep_data['avg_list_items']} items each) — INCLUDE at least {max(int(deep_data['avg_list_count']), 4)} lists\n"
            f"  - FAQ items: {deep_data['avg_faq_count']} (max {deep_data['max_faq_count']}) — generate at least {max(int(deep_data['avg_faq_count']), 7)} FAQ\n"
            f"  - Schema blocks: {deep_data['avg_schema_blocks']} (we generate JSON-LD automatically)\n\n"

            + (f"▸ KEYWORD PLACEMENT (competitors do this):\n"
               f"  - KW density: {deep_data['avg_kw_density_pct']}% — match this rate\n"
               f"  - KW in H1: {deep_data['kw_in_h1_pct']:.0f}% of competitors include it — MUST include in yours\n"
               f"  - KW in H2: avg {deep_data['avg_kw_in_h2']}x — use in {max(int(deep_data['avg_kw_in_h2']), 2)} H2 headings\n"
               f"  - KW in H3: avg {deep_data['avg_kw_in_h3']}x\n"
               f"  - KW in first 100 words: {deep_data['kw_in_first100_pct']:.0f}% of competitors do this — MUST do same\n"
               f"  - KW in last 100 words (conclusion): {deep_data['kw_in_last100_pct']:.0f}% — include in conclusion\n"
               f"  - KW in image alt texts: avg {deep_data['avg_kw_in_alt']}x\n"
               f"  - KW in anchor texts: avg {deep_data['avg_kw_in_anchors']}x\n\n"
               if deep_data.get('avg_kw_count') else "")

            + (f"▸ LSI / SEMANTIC TERMS (top words competitors use — INTEGRATE these naturally):\n"
               f"  - Unique terms: {deep_data['avg_lsi_unique']} — use varied vocabulary\n"
               f"  - TOP 15 LSI terms (with frequencies): {top_lsi_str}\n"
               f"  - You MUST use the top 10-15 of these terms throughout your text — they signal topic depth to Google\n\n"
               if top_lsi_str else "")

            + (f"▸ LONG-TAIL PHRASES (use 5–10 of these patterns):\n"
               f"  Examples from competitors: {long_tail_str}\n\n"
               if long_tail_str else "")

            + (f"▸ LINKS STRATEGY:\n"
               f"  - Internal links: {deep_data['avg_internal_links']} avg (we add later via post-processing)\n"
               f"  - External links: {deep_data['avg_external_links']} avg to {deep_data['avg_unique_ext_domains']} unique domains\n"
               f"  - Link density: {deep_data['avg_links_per_1000_words']} per 1000 words\n"
               f"  - Avg anchor length: {deep_data['avg_anchor_words']} words\n"
               + (f"  - COMMON EXTERNAL SOURCES competitors link to (cite these for E-E-A-T): {common_dom_str}\n" if common_dom_str else "")
               + "\n")

            + (f"▸ ON-PAGE SEO TARGETS:\n"
               f"  - Title length: {deep_data['avg_title_length']} chars (sweet spot 50–60)\n"
               f"  - Meta Description: {deep_data['avg_meta_desc_length']} chars (sweet spot 130–160)\n"
               + (f"  - Sample competitor titles: {sample_titles}\n" if sample_titles else "")
               + "\n")

            + (f"▸ READABILITY (compete with these):\n"
               f"  - Avg sentence length: {deep_data['avg_sentence_len']} words — your target 8–21\n"
               f"  - Long sentences (>30w): {deep_data['avg_long_sentence_pct']}% — keep yours under 10%\n\n")

            + f"⚠️ COMPETE PRINCIPLE: match or BEAT each metric. Use top LSI terms naturally, include all H2-relevant KWs, build comparable structure depth.\n\n"
        )

    # AGGRESSIVE PROFILE — based on analysis of 6 top-ranking sites
    aggressive_rule = (
        "═══ 🔥 AGGRESSIVE PROFILE — MATCH TOP-RANKING SITES ═══\n"
        "⚠️ These targets are based on analysis of 6 top-ranking Italian casino SEO sites.\n"
        "You MUST hit these numbers or the text fails QA.\n\n"
        "▸ BODY WORDS: target 3500–4000 слов.\n"
        f"▸ H2 HEADINGS: 18–22 sections (minimum 18). Top sites average 19.7 H2.\n"
        f"▸ H3 HEADINGS: 35–45 subsections (minimum 35). Top sites average 40.5 H3.\n"
        f"▸ TABLES: 6–8 tables, each with 6–10 rows. Top sites average 7.8 tables × 8.5 rows.\n"
        f"▸ LISTS: 30–40 bullet/numbered lists, each with 4–6 items. Top sites average 37 lists.\n"
        f"▸ PARAGRAPH LENGTH: 40–60 words per paragraph (top sites average 52w, max 150w).\n"
        f"▸ LSI RICHNESS: use 1500+ UNIQUE vocabulary words — do not repeat same phrases.\n"
        f"▸ TERMINOLOGY VARIATIONS — use SYNONYMS aggressively:\n"
        f"    - For 'casino non AAMS' ALSO use: 'casinò esteri', 'casino stranieri', 'casino senza licenza ADM', 'piattaforme estere', 'siti esteri', 'casino internazionali'\n"
        f"    - For 'bonus': 'promozione', 'offerta', 'vantaggio', 'incentivo', 'omaggio'\n"
        f"    - For 'giocatori': 'utenti', 'clienti', 'giocatore italiano', 'appassionati', 'scommettitori'\n"
        f"    - For 'gioco': 'intrattenimento', 'scommessa', 'puntata', 'partita'\n"
        f"▸ TOP LSI TERMS to use repeatedly (aim for these frequencies):\n"
        f"    bonus (50–100x), casinò (40–90x), giochi (30–60x), fino (25–55x), gioco (25–55x),\n"
        f"    giocatori (25–50x), stranieri (20–45x), licenza (20–40x), online (20–40x),\n"
        f"    giri (15–35x), slot (15–35x), live (15–30x), esteri (15–30x),\n"
        f"    piattaforme (10–25x), deposito (10–25x), migliori (10–25x),\n"
        f"    promozioni (8–20x), senza (8–20x), benvenuto (5–15x), spin (5–15x)\n"
        f"▸ NO FORCED FAQ section — top sites DON'T have FAQ (0 FAQ on all 6 sites).\n"
        f"    Skip FAQ entirely OR replace with deeper topic H2 sections.\n"
        f"▸ NO external links — top sites DON'T link out (0 external links).\n"
        f"▸ LONG-TAIL: generate 150+ unique long-tail phrases naturally (3–5 word combinations).\n"
        f"▸ TITLE length: ~62 chars (sweet spot for this niche).\n"
        f"▸ META length: ~100 chars (top sites use shorter meta descriptions).\n"
        f"▸ SENTENCE LENGTH: avg 22–28 words (top sites use longer, more detailed sentences).\n\n"
        f"⚠️ CRITICAL MINDSET: Do not generate generic 4000-word review — write a COMPREHENSIVE,\n"
        f"DEEP authoritative guide that covers EVERY aspect of the topic from multiple angles.\n"
        f"Each H2 section must be 300–500 words with 2–4 H3 subsections inside.\n\n"
    )

    # Dynamic word count from SERP targets
    if serp_targets and serp_targets.get("word_count"):
        _wc_tup = serp_targets["word_count"]
        _v3_wc_min = int(_wc_tup[0])
        _v3_wc_max = int(_wc_tup[1])
        _v3_wc_target = (_v3_wc_min + _v3_wc_max) // 2
        _v3_wc_n = serp_targets.get("_n_competitors", 3)
    else:
        _v3_wc_target = word_count
        _v3_wc_min = int(word_count * 0.9)
        _v3_wc_max = int(word_count * 1.15)
        _v3_wc_n = 0
    base_rules = tf_rule + struct_rule + deep_rule + aggressive_rule + (
        "═══════════════════════════════════════════\n"
        "MANDATORY SEO & QUALITY RULES\n"
        "═══════════════════════════════════════════\n\n"
        f"▸ WORD COUNT: {_v3_wc_target} слов (на основе анализа топ-{_v3_wc_n} конкурентов в выдаче). Допустимый диапазон: {_v3_wc_min}–{_v3_wc_max} слов. Закончи текст естественно, не обрывай мысль.\n\n"
        f"🚨🚨🚨 LANGUAGE LOCK — АБСОЛЮТНОЕ ПРАВИЛО № 1:\n"
        f"  🔴 ВЕСЬ текст ОБЯЗАТЕЛЬНО НА ЯЗЫКЕ: {lang}\n"
        f"  🔴 НИ ОДНОГО СЛОВА, ПРЕДЛОЖЕНИЯ ИЛИ ФРАЗЫ НА ДРУГОМ ЯЗЫКЕ\n"
        f"  🔴 ЗАПРЕЩЕНО смешивать языки в одном заголовке, предложении, параграфе\n\n"
        f"  ПРИМЕР БРАКА (ТАК НЕ ПИСАТЬ):\n"
        f"    H2: Mobile Casino Experience — Aplikacja vs Przeglądarka  ❌ (mix EN+PL)\n"
        f"    H2: Top 5 Casino Italy 2026 — Migliori Scelte  ❌ (mix EN+IT)\n"
        f"    H2: Bonuses & Promotions für Spieler  ❌ (mix EN+DE)\n"
        f"    H2: Payment Methods в Kasynach  ❌ (mix EN+PL+RU)\n"
        f"  ПРАВИЛЬНО ДЛЯ {lang} (ПИШИ ТАК):\n"
        + (f"    H2: Mobilne kasyno — Aplikacja vs Przeglądarka ✅\n"
           f"    H2: Top 5 Kasyn Online Polska 2026 ✅\n"
           f"    H2: Bonusy i promocje dla graczy ✅\n" if lang == 'PL' else
           f"    H2: Esperienza Mobile — App vs Browser ✅\n"
           f"    H2: Top 5 Casinò Italia 2026 ✅\n"
           f"    H2: Bonus e Promozioni per i Giocatori ✅\n" if lang == 'IT' else
           f"    H2: Mobiles Casino — App vs Browser ✅\n"
           f"    H2: Top 5 Online Casinos Deutschland 2026 ✅\n"
           f"    H2: Bonusse & Promotionen für Spieler ✅\n" if lang == 'DE' else
           f"    H2: Mobiel Casino — App vs Browser ✅\n"
           f"    H2: Top 5 Online Casino's Nederland 2026 ✅\n"
           f"    H2: Bonussen en Promoties voor Spelers ✅\n" if lang == 'NL' else
           f"    H2: Casino Móvil — App vs Navegador ✅\n"
           f"    H2: Top 5 Casinos Online España 2026 ✅\n" if lang == 'ES' else
           f"    ВСЕ заголовки только на {lang}\n")
        + f"\n  ЧЕК-ЛИСТ ПЕРЕД ОТПРАВКОЙ ТЕКСТА:\n"
        f"    ☑ Все H1/H2/H3/H4 в тексте — на {lang}?\n"
        f"    ☑ Все заголовки таблиц — на {lang}?\n"
        f"    ☑ Все ячейки таблиц — на {lang}? (кроме брендов, валют, %)\n"
        f"    ☑ Все вопросы FAQ — на {lang}?\n"
        f"    ☑ Все CTA кнопки — на {lang}? (не 'Play Now', а на {lang})\n"
        f"    ☑ Все bullet пункты — на {lang}?\n\n"
        f"  Исключения (ЭТО ОК):\n"
        f"    ✅ Названия брендов: SpinLander, Wild Tokyo, Bet365, Slotuna\n"
        f"    ✅ Регуляторы: ADM, MGA, KSA, KSA.nl, ESBK\n"
        f"    ✅ Валюты: EUR, USD, GBP, PLN, HUF\n"
        f"    ✅ Технические бренды: PayPal, Skrill, Visa, Mastercard, Bitcoin, iDEAL\n"
        f"    ✅ URL и email адреса\n\n"
        f"  ⚠️ Текст с примесью другого языка (даже 1 заголовок) — ЭТО БРАК.\n"
        f"  ⚠️ БОТ АВТОМАТИЧЕСКИ ОТКЛОНИТ И ПОВТОРИТ ГЕНЕРАЦИЮ.\n\n"
        "▸ HEADINGS — ABSOLUTE RULE (no exceptions):\n"
        "  EVERY heading in the article MUST use the prefix format:\n"
        "  H1: [text]   H2: [text]   H3: [text]   H4: [text]\n"
        "  FORBIDDEN: # Heading, ## Heading, **Heading**, HEADING IN CAPS\n"
        "  If any heading is missing its prefix → the entire text fails QA\n"
        "  This applies to ALL sections without exception: intro, ranking, reviews, FAQ, conclusion\n\n"
        "▸ READABILITY TARGETS (CRITICAL):\n"
        f"  - LANGUAGE: {lang}\n"
        + ("  - FKGL target: Grade 7–9 (English — 0.39×words/sentences + 11.8×syllables/words - 15.59)\n" if lang.startswith("en") else "")
        + "  - LIX target: <45 (universal — A/B + C×100/A where A=words, B=sentences, C=words>6 letters)\n"
        + "  - If LIX >55 = too complex — use shorter sentences, simpler vocabulary\n\n"
        "▸ SENTENCE LENGTH (readability rule — CRITICAL):\n"
        "  - OPTIMAL range: 8–21 words per sentence\n"
        "  - MINIMUM: ≥8 words (shorter = primitive, choppy)\n"
        "  - MAXIMUM: ≤21 words (longer = hard to read, avoid)\n"
        "  - MIX lengths within 8–21 range for natural rhythm\n"
        "  - Target: 70%+ of sentences in optimal range\n\n"
        "▸ PARAGRAPHS — STRICT RULE:\n"
        "  - MAXIMUM 3 sentences per paragraph. This is a hard limit — never exceed.\n"
        "  - MAXIMUM 60 words per paragraph. If longer — you MUST split into 2 paragraphs.\n"
        "  - Each paragraph = ONE single idea only\n"
        "  - Separate every paragraph with empty line\n"
        "  - NEVER write walls of text. White space must appear every 3 sentences.\n\n"
        f"▸ KEYWORD DENSITY — PRIMARY KEY '{primary_kw}':\n"
        f"  - Use '{primary_kw}' {kw_min}–{kw_max} times across the full article\n"
        f"  - DENSITY FORMULA: 1–3 mentions per every 100 words = {kw_min}–{kw_max} total for target length\n"
        f"  - PRACTICAL RULE: per 1000 characters (~150–200 words) → use keyword 3–6 times\n"
        "  - Place in: H1, first paragraph, 2–3 H2 headings, conclusion, FAQ\n"
        "  - Do NOT cluster — spread EVENLY: every section gets 1–2 mentions\n"
        "  - NEVER repeat same phrase in same sentence or adjacent sentences\n"
        "  - NEVER stuff: 3+ mentions in one paragraph = keyword stuffing, penalized\n"
        + (f"\n▸ SECONDARY KEYWORDS (use each 1–3 times naturally): {', '.join(secondary_kws)}\n" if secondary_kws else "") +
        "\n▸ ENTITY COVERAGE — include these in context where relevant:\n"
        "  PRIMARY entities: casino name, country/geo, license authority, regulator\n"
        "  TRUST entities: license number/type, payment methods, withdrawal timing,\n"
        "    KYC/verification process, customer support channels, responsible gambling tools\n"
        "  RELATED entities: game providers, bonus types, wagering requirements, currencies\n\n"
        "▸ TRUST BLOCKS — include where relevant to the topic:\n"
        "  - Licensing & regulatory status (which authority, what it means for player)\n"
        "  - Payment clarity (methods, limits, processing time, fees)\n"
        "  - KYC / account verification requirements\n"
        "  - Responsible gambling (brief, not preachy — tools available, limits, self-exclusion)\n"
        "  - Realistic limitations & caveats (geo restrictions, max bonus, excluded countries)\n\n"
        "▸ STYLE & VOICE:\n"
        "  - Expert tone: confident, direct, data-backed\n"
        "  - Active voice throughout\n"
        "  - NO filler: 'it is worth noting', 'needless to say', 'in conclusion we can say'\n"
        "  - Include specific numbers: percentages, amounts, timeframes, ratings\n"
        "  - Avoid glamorizing losses or risky gambling behavior\n\n"
        "▸ STRUCTURE REQUIREMENTS:\n"
        "  - MINIMUM 18 H2 sections, each 300–500 words (aggressive profile)\n"
        "  - MINIMUM 35 H3 subsections across all H2 blocks\n"
        "  - H3 subsections inside major H2 blocks\n"
        "  - ОБЯЗАТЕЛЬНО минимум 10 таблиц в тексте (НЕ МЕНЕЕ 10!). Каждая секция про платёжные методы, бонусы, сравнения, лицензии — ТОЛЬКО в виде таблицы.\n"
        "    Каждая таблица: минимум 3 колонки и 5 строк данных (не считая заголовка).\n"
        "    Обязательные типы таблиц: (1) Сравнение казино (Бренд|Бонус|Лицензия|Выплаты|Рейтинг), (2) Платёжные методы (Метод|Депозит|Вывод|Комиссия|Мин.сумма), (3) Бонусы (Тип|Сумма|Вейджер|Срок действия|Игры), (4) Провайдеры игр (Провайдер|Слоты|Live|Джекпот|RTP), (5) Лицензии и регуляторы (Юрисдикция|Регулятор|Гарантии|Проверки), (6) Мобильная поддержка (Платформа|iOS|Android|Браузер|Приложение), (7) Лимиты депозита/вывода (Метод|Мин|Макс|Комиссия|Срок). Добавь ещё 3+ таблиц по теме.\n"
        "  - Minimum 2 bullet or numbered lists\n\n"
        "▸ END ORDER (mandatory):\n"
        "  Conclusion (200+ words summarizing key points) →\n"
        "  FAQ LAST: ОБЯЗАТЕЛЬНО минимум 16 вопросов с развёрнутыми ответами (3-5 предложений = 50-100 слов каждый). Вопросы должны покрывать все ключевые слова темы.\n"
        "  FAQ must cover: legal status, bonus claim, withdrawals, safety, mobile, min deposit, support\n"
        "  FAQ FORMAT — STRICT: each question as plain bold heading (e.g. How do I claim a bonus?), then answer paragraph.\n"
        "  NEVER use Q:, A:, Q1., 1), letters (a), b)) or any prefix labels before questions or answers.\n"
        "  NEVER use bullet points inside FAQ answers. Plain paragraphs only.\n\n"
        "═══ E-E-A-T WRITING QUALITY ═══\n\n"
        f"▸ CURRENT YEAR — CRITICAL: The year is 2026. Always write '2026' in titles, headers, and throughout the text. NEVER write 2024 or 2025 as the current year. All statistics, ratings, and data should reference 2026.\n\n"
        "▸ INTRO: answer 3 questions in first 2 sentences:\n"
        "  1. What is this page? 2. Who is it for? 3. What value does reader get?\n"
        "  GOOD: 'UK players in 2026 have 35+ licensed casinos — this guide ranks the top 5 by bonus value, payout speed, and game range.'\n\n"
        "▸ SPECIFICITY — every claim needs a concrete fact:\n"
        "  - Bonus: exact % + cap + wagering + expiry + eligible games\n"
        "  - Payments: exact method name + processing time + fees + limits\n"
        "  - Licensing: exact authority + what it guarantees for the player\n"
        "  NEVER write: 'The casino offers generous bonuses and fast withdrawals.'\n"
        "  WRITE: 'The 100%/£200 bonus has 35x wagering on bonus only. PayPal withdrawals take 4–12 hours.'\n\n"
        "▸ EVIDENCE — back every key claim with proof or comparison:\n"
        "  - 'Players report...', 'Testing showed...', 'License authority requires...'\n"
        "  - Use comparisons: 'faster than the 3-day industry average'\n"
        "  - Each H2 section: minimum 1 evidence-backed statement\n\n"
        "▸ STRUCTURE — use lists, steps, tables instead of prose:\n"
        "  - Registration → numbered steps (NEVER describe in prose)\n"
        "  - Bonus comparison → table: Casino | Bonus | Wagering | Expiry\n"
        "  - Pros/Cons → bullet two-column list\n"
        "  - Payments → table or bullets with exact specs\n\n"
        "▸ HEADINGS — max 7 words, specific, no vague nouns:\n"
        "  BAD: 'Information About Bonuses' → GOOD: 'Welcome Bonus: 100% to £200 — Full Terms'\n"
        "  BAD: 'Mobile Gaming Overview' → GOOD: 'Casino App: iOS & Android Download Guide'\n\n"
        "▸ SENTENCE DISCIPLINE — no water, minimize pronouns:\n"
        "  - Every sentence = one complete thought + one fact\n"
        "  - Avoid sentence starters: 'We', 'They', 'It', 'This', 'That'\n"
        "  - No filler: 'It should be noted', 'As we can see', 'In fact'\n"
        "  - Active voice: 'Players receive 200 FS' not '200 FS are received'\n"
        "  - If deleting a sentence changes nothing → delete it"
        "\n▸ PERSONAL EXPERIENCE — CRITICAL for E-E-A-T:\n"
        "  - Write as a real expert who has TESTED the casino: 'During our test, withdrawal took 4h via PayPal', 'We registered in 3 minutes — no ID required initially'\n"
        "  - Add specific observations: 'The live chat responded in 47 seconds', 'Bonus terms are buried in FAQ section 5 — took time to find'\n"
        "  - Compare to real experience: 'Unlike most casinos, this one shows wagering progress in the lobby'\n"
        "  - Each casino H3 block must include at least 1 personal experience sentence ('When we tested...', 'Our review found...', 'We noticed...')\n"
        "  - NEVER write generic praise: 'This is a great casino.' — always back with a specific observation\n\n"
    )

    req_line = f"Additional requirements: {requirements}" if requirements else ""

    # Build SERP targets prompt block (extended with TF-IDF, Topic Map, Featured Snippet)
    if serp_targets:
        _serp_htmls = serp_targets.get("_serp_htmls", [])
        _serp_lang = task.get("language", "en").lower()[:2]
        _serp_geo = task.get("geo", "us")
        _serp_block = build_extended_serp_prompt_block(
            serp_targets,
            primary_kw,
            competitor_htmls=_serp_htmls if _serp_htmls else None,
            lang=_serp_lang,
            geo=_serp_geo,
        )
    else:
        _serp_block = ""

    # Используем предварительно сгенерированный уникальный title если есть
    _pre_meta_title = ""
    _pre_h1 = ""
    _pre_meta_desc = ""
    if pre_title and isinstance(pre_title, dict):
        _pre_meta_title = pre_title.get("meta_title", "")
        _pre_h1 = pre_title.get("h1", "")
        _pre_meta_desc = pre_title.get("meta_desc", "")

    _title_instruction = (
        f"Meta Title: USE EXACTLY THIS TITLE (do not change): {_pre_meta_title}\n"
        f"Meta Description: USE EXACTLY THIS DESCRIPTION (do not change): {_pre_meta_desc}\n"
        if _pre_meta_title else
        f"Meta Title: [SEO title 50-60 chars, include primary keyword + {current_year}, MUST BE UNIQUE from other texts in batch]\n"
        f"Meta Description: [description 130-155 chars, include primary keyword, clear value prop, MUST BE UNIQUE]\n"
    )
    _h1_instruction = (
        f"H1: USE EXACTLY THIS H1: {_pre_h1}" if _pre_h1 else
        f"H1: [unique heading for THIS text only — must differ from all other texts in batch]"
    )

    # text_mono — deep single brand review (landing page structure)
    if task_type == "text_mono" and casino:
        return _serp_block + (
            f"Write a comprehensive, expert-level single-brand SEO article about {casino} for players in {geo}.\n\n"
            f"Language: {lang}\n"
            f"Target keywords to integrate naturally: {kw_str}\n"
            f"{req_line}\n\n"
            f"OUTPUT FORMAT — START THE DOCUMENT WITH:\n"
            f"{_title_instruction}"
            f"(leave one empty line after Meta Description, then start article)\n\n"
            f"🚫🚫🚫 SINGLE-BRAND MONO ARTICLE — ABSOLUTE RULES:\n"
            f"1. THIS IS A MONO (single-brand) ARTICLE ABOUT {casino} ONLY.\n"
            f"2. NEVER mention, compare, or list any other casino brands. No competitor names, no 'Casino X', no alternative casinos.\n"
            f"3. NEVER create comparison tables with multiple casinos (Casino | Rating | Bonus columns = FORBIDDEN).\n"
            f"4. ALL tables must be about {casino} features only (e.g. payment methods, game categories, bonus terms).\n"
            f"5. VIOLATION = article rejected. Keep 100% focus on {casino}.\n\n"
            f"🚨 CRITICAL: ALL H2/H3 headings below are DESCRIPTIONS of section topics in English for your understanding.\n"
            f"YOU MUST WRITE ALL ACTUAL HEADINGS IN {lang} LANGUAGE. Do NOT copy English headings into the article.\n\n"
            f"MANDATORY STRUCTURE (in this exact order):\n"
            f"[H1 in {lang}]: about {casino} — main value proposition in {geo}\n"
            f"[Intro in {lang}] (60 words max): what platform is, for whom, key USP + main keyword\n\n"
            f"[H2 in {lang}]: Why choose {casino}?\n"
            f"  [H3 in {lang}]: Top odds / games selection\n"
            f"  [H3 in {lang}]: Bonuses & promotions\n"
            f"  [H3 in {lang}]: Security & licensing\n"
            f"  [H3 in {lang}]: Mobile experience\n"
            f"  ⚠️ NO comparison table with other casinos here. Only {casino} features.\n\n"
            f"[H2 in {lang}]: How to get started at {casino}\n"
            f"  Numbered steps: Sign up → Verify → Deposit → Claim bonus → Start playing (all in {lang})\n\n"
            f"[H2 in {lang}]: {casino} app & mobile platform\n"
            f"  Android/iOS availability, features, how to download\n\n"
            f"[H2 in {lang}]: Games / sports / products available at {casino}\n"
            f"  Detailed categories with descriptions. Table allowed ONLY for game categories (Slots | Live | Table — all {casino} content).\n"
            f"  ⚠️ Do NOT list other casino brands in this table.\n\n"
            f"[H2 in {lang}]: Bonuses & promotions at {casino}\n"
            f"  Welcome bonus, ongoing promos, loyalty program, T&Cs\n\n"
            f"[H2 in {lang}]: Deposits & withdrawals\n"
            f"  Payment methods table (method | min | max | time) — {casino} only\n\n"
            f"[H2 in {lang}]: {casino} login & account security\n"
            f"  Login steps, SSL, 2FA, account verification\n\n"
            f"[H2 in {lang}]: FAQ (minimum 6 questions specific to {casino})\n"
            f"  FORMAT STRICT: question as plain bold line, answer as paragraph. NO Q:, A:, Q1., 1), a), b) labels ever.\n"
            f"  ALL questions and answers in {lang}.\n\n"
            f"❗ REMINDER: The actual H2/H3 text in the output MUST BE IN {lang}, not English.\n"
            f"Example for Polish: 'H2: Dlaczego warto wybrać {casino}?' (NOT 'H2: Why Choose {casino}?')\n"
            f"Example for Italian: 'H2: Perché scegliere {casino}?'\n"
            f"Example for German: 'H2: Warum {casino} wählen?'\n\n"
            f"{base_rules}"
        )

    # text_review — multi-brand comparison article
    elif task_type == "text_review" or not casino:
        kw_primary = keywords[0] if keywords else f"best casinos {geo}"
        kw_secondary = ", ".join(keywords[1:4]) if len(keywords) > 1 else ""
        return (
            f"Write a comprehensive, expert-level SEO article about the best online casinos in {geo}.\n\n"
            f"Language: {lang}\n"
            f"PRIMARY keyword (use in H1, first paragraph, 2–3 H2s): {kw_primary}\n"
            f"SECONDARY keywords (use naturally throughout, 1–2x each): {kw_secondary}\n"
            f"All keywords to integrate: {kw_str}\n"
            f"BRANDS RULE:\n"
            + (f"  - MANDATORY BRANDS (from task, must be in EVERY text): {casinos_str}\n"
               f"  - MINIMUM 10 brands per text: if task has fewer than 10, find additional relevant {geo} market brands yourself and add them\n"
               f"  - ALL brands (task brands + your additions) must be described in tables and short descriptions\n"
               f"  - EVERY text must contain: {casinos_str} + your own additions up to 10 total\n"
               if brands_from_task and casinos_str else
               f"  - No brands specified — extract brand names from competitor pages and use those\n"
               f"  - Create a realistic Top {len(all_casinos) if all_casinos else 5} list based on competitor research\n")
            + f"Casinos to feature and compare: {casinos_str if casinos_str else '[extract from competitor research]'}\n"
            f"{req_line}\n\n"
            f"OUTPUT FORMAT — START THE DOCUMENT WITH:\n"
            f"{_title_instruction}"
            f"(leave one empty line after Meta Description, then start article)\n\n"
            f"THIS IS TEXT #{text_index + 1} IN THE BATCH. Each text must differ from others in structure and brand descriptions.\n\n"
            f"STRUCTURE ORDER FOR THIS TEXT (follow exactly, do NOT reorder):\n"
            f"{struct_order}\n\n"
            f"BRAND DESCRIPTION RULE FOR THIS TEXT:\n"
            f"{brand_desc_rule}\n\n"
            f"🚨🚨🚨 АБСОЛЮТНО ОБЯЗАТЕЛЬНОЕ ПРАВИЛО о H1:\n"
            f"КАЖДЫЙ ТЕКСТ ОБЯЗАТЕЛЬНО НАЧИНАЕТСЯ С 'H1: ' (с префиксом 'H1:' и пробелом!).\n"
            f"ТЕКСТ БЕЗ H1 — ЭТО БРАК. БОТ ОТКЛОНИТ ТЕГР.\n"
            f"ПОРЯДОК СТРОГО:\n"
            f"  СТРОКА 1: Meta Title: ...\n"
            f"  СТРОКА 2: Meta Description: ...\n"
            f"  СТРОКА 3: (пустая)\n"
            f"  СТРОКА 4: H1: [твой заголовок на {lang}]\n"
            f"  СТРОКА 5: (пустая)\n"
            f"  СТРОКА 6: [вступительный параграф]\n"
            f"  СТРОКА 7+: H2: ... далее по структуре\n"
            f"ЗАПРЕЩЕНО:\n"
            f"  - Начинать сразу с H2 (без H1)\n"
            f"  - Вставлять заголовок без префикса 'H1:'\n"
            f"  - Использовать markdown '# Заголовок' вместо 'H1: Заголовок'\n"
            f"  - Текст без H1 вообще\n\n"
            f"MANDATORY ARTICLE STRUCTURE — SNIPPET-FIRST:\n"
            f"=== PAGE STRUCTURE (follow exactly) ===\n"
            f"H1 RULE — make it natural and varied (mirror competitor H1 styles):\n"
            f"  - FOR THIS TEXT USE H1 POSITION: {h1_position}\n"
            f"  - POSITION START: [keyword] {geo} {current_year} — [natural phrase in {lang}]\n"
            f"    Example: Online Casino Hrvatska 2026 – Najbolji izbori i savjeti\n"
            f"  - POSITION MIDDLE: [natural phrase] [keyword] {geo} {current_year}\n"
            f"    Example: Najbolji {geo} [keyword] 2026 — Vodič za početnike\n"
            f"  - POSITION END: [natural phrase] {geo} {current_year} — [keyword]\n"
            f"    Example: Igrajte sigurno u {geo} 2026 — [keyword] pregled\n"
            f"  - Always include: ВЧ keyword + GEO ({geo}) + year ({current_year})\n"
            f"  - H1 must feel natural in {lang}, not mechanical\n"
            f"  - CRITICAL: H1 must be in {lang} language — NEVER in English for non-English articles\n"
            f"  - CRITICAL: Do NOT write 'Expert Guide' or '[keyword] — 2026 Expert Guide' — this is forbidden\n"
            f"  - Use natural phrasing in {lang}: Vodič, Przewodnik, Guida, Leitfaden, Gids, etc.\n"
            f"  - Make H1 sound natural in {lang}, not a keyword dump\n\n"
            f"INTRO PARAGRAPH (REQUIRED — 3-4 sentences, BEFORE first H2):\n"
            f"  Write a concise intro paragraph AFTER H1 and BEFORE the first H2.\n"
            f"  - Answer: who is this for, what problem does it solve, what will reader find\n"
            f"  - Include primary keyword naturally\n"
            f"  - Max 60 words, no fluff, no brand names, no history\n\n"
            f"CORRECT OUTPUT FORMAT (write article in this order):\n"
            f"  1. H1: [your natural H1 per rules above]\n"
            f"  2. [Intro paragraph — 3-4 sentences]\n"
            f"  3. H2: [snippet H2 with keyword+GEO+year]\n"
            f"  FORMAT: Top {max(5, len(all_casinos) if all_casinos else 5)} [primary keyword] {geo} {current_year}\n"
            f"  OR: Best [primary keyword] in {geo} ({current_year})\n"
            f"  MUST CONTAIN: keyword + GEO ({geo}) + year ({current_year}) + number (Top N)\n"
            f"  EXAMPLE H2: Top 5 Wypłacalne Kasyna Online Polska {current_year}\n\n"
            "PARAGRAPH UNDER FIRST H2 — 2-3 sentences ONLY:\n"
            "  Sentence 1: Direct answer to user intent — explain what this list is\n"
            "  Sentence 2: Selection criteria (bonuses, license, payout speed, games, mobile)\n"
            "  Sentence 3: Lead into comparison table or Top list below\n"
            "  FORBIDDEN: long intros, brand history, vague phrases like many players prefer\n\n"
            "COMPARISON TABLE — MANDATORY after intro paragraph:\n"
            f"  FORMAT: pipe table — translate column headers to {lang}:\n"
            f"  EN: Brand | Bonus | Key Strengths | Payments | Speed | Rating\n"
            f"  NL: Brand | Bonus | Voordelen | Betaalmethoden | Snelheid | Score\n"
            f"  PL: Marka | Bonus | Zalety | Płatności | Szybkość | Ocena\n"
            f"  HR/BS: Brand | Bonus | Prednosti | Plaćanja | Brzina | Ocjena\n"
            f"  DE: Marke | Bonus | Vorteile | Zahlungen | Geschwindigkeit | Bewertung\n"
            f"  IT: Brand | Bonus | Punti Forti | Pagamenti | Velocità | Valutazione\n"
            f"  PT: Marca | Bônus | Pontos Fortes | Pagamentos | Velocidade | Nota\n"
            f"  Use language {lang} for ALL column headers — never mix languages\n\n"
            "  TABLE HEADING RULE: Do NOT use descriptive heading like Comparison Table or Usporedna tablica\n"
            "  Use a short natural H2 heading that includes primary keyword, like: Top N Kasyna — Rangirane 2026\n"
            "  The table must have at least 5 rows (one per casino) and all columns filled\n"
            "  Place table IMMEDIATELY after the 2-3 sentence intro paragraph\n\n"
            "FORBIDDEN PATTERNS:\n"
            "  X Long intro before first useful block\n"
            "  X Abstract phrases: there are many options, in todays market\n"
            "  X Starting page with brand history\n"
            "  X Keyword stuffing\n\n"
            f"--- Now write the article following the STRUCTURE ORDER defined above ---\n\n"
            f"RANKING SECTION FORMAT (when you write the ranking H2):\n"
            f"  H3: Casino Name (Rating: X.X/10) — use H3: prefix\n"
            f"  Welcome bonus: [details] — plain paragraph\n"
            f"  3 key strengths (bullets), 1-2 weaknesses (bullets)\n"
            f"  Payment methods: [list] — plain paragraph\n"
            f"  Best for: [player type] — plain paragraph\n"
            f"  RULE: ONLY strengths/weaknesses use - bullets. All other fields = plain paragraphs.\n\n"
            f"SECTION DETAILS (for each section in your structure order above):\n"
            f"  - Ranking/Top list: include comparison table + H3 per brand (use {brand_style} description style)\n"
            f"  - Bonuses: types of bonuses, wagering requirements, tips for {geo} players\n"
            f"  - Payments: pipe table Method|Deposit time|Withdrawal|Fees|Min amount\n"
            f"  - Mobile: app vs browser, best mobile picks, responsive design\n"
            f"  - Responsible: self-exclusion tools, deposit limits, support resources\n"
            f"  - Conclusion: 200+ words summary + final top picks\n"
            f"  - FAQ: minimum 7 questions, 50-100 words each, no Q:/A: labels\n\n"
            f"{base_rules}"
        )

    # Legacy: single casino review (old format)
    else:
        return (
            f"Write a detailed SEO casino review for {casino} targeting {geo} players.\n\n"
            f"Language: {lang}\n"
            f"Target keywords: {kw_str}\n"
            f"{req_line}\n\n"
            f"Sections to cover:\n"
            f"- Overview & Licensing\n"
            f"- Welcome Bonus & Promotions\n"
            f"- Games & Software Providers\n"
            f"- Payment Methods & Withdrawals\n"
            f"- Mobile Experience\n"
            f"- Customer Support\n"
            f"- Pros & Cons table\n"
            f"- Conclusion\n"
            f"- FAQ (5+ questions)\n\n"
            f"{base_rules}"
        )


# ─── Main ─────────────────────────────────────────────────────────────────────




# ══════════════════════════════════════════════════════════════════════════════
# ▼▼▼  v3: AHREFS-DRIVEN TARGETS + CONTENT AUDIT  ▼▼▼
# ══════════════════════════════════════════════════════════════════════════════

import math as _math_v3
try:
    from bs4 import BeautifulSoup as _BS
    _BS_OK = True
except ImportError:
    _BS_OK = False

# ── Google-guideline hard caps (applied on top of SERP medians) ──
_KW_DENSITY_MAX   = 3.0   # %
_LONG_SENT_MAX    = 5.0   # % sentences >30 words
_FLESCH_MIN       = 10.0  # readability floor


def _flesch_ease(text: str) -> float:
    """Approximation of Flesch Reading Ease for any language."""
    words = re.findall(r'\b\w+\b', text)
    sentences = re.split(r'[.!?]+', text)
    sentences = [s.strip() for s in sentences if s.strip()]
    syllables = sum(_count_syllables(w) for w in words)
    if not sentences or not words:
        return 0.0
    asl = len(words) / len(sentences)  # avg sentence length
    asw = syllables / len(words)        # avg syllables per word
    score = 206.835 - 1.015 * asl - 84.6 * asw
    return round(score, 1)


def _count_syllables(word: str) -> int:
    """Naive syllable counter (vowel groups)."""
    word = word.lower()
    count = len(re.findall(r'[aeiouáéíóúàèìòùäëïöüāēīōū]+', word))
    return max(1, count)


def _extract_metrics_from_html(html: str, keyword: str = "") -> dict:
    """
    Extract ALL content metrics from raw HTML or plain text.
    Returns dict with all measurable values.
    """
    kw_lower = keyword.lower().strip()

    if _BS_OK:
        soup = _BS(html, "html.parser")
        # Remove script/style
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()
        plain = soup.get_text(separator=" ")
        all_text = plain
        h1_tags  = [t.get_text() for t in soup.find_all("h1")]
        h2_tags  = [t.get_text() for t in soup.find_all("h2")]
        h3_tags  = [t.get_text() for t in soup.find_all("h3")]
        h4_tags  = [t.get_text() for t in soup.find_all("h4")]
        tables   = soup.find_all("table")
        lists_ul = soup.find_all("ul")
        lists_ol = soup.find_all("ol")
        anchors  = [a.get_text() for a in soup.find_all("a")]
        # FAQ detection: look for dl/details or divs with FAQ patterns
        faq_count = len(soup.find_all(["details"])) + \
                    len([t for t in soup.find_all(string=re.compile(r'^\s*(Q:|A:|FAQ)', re.I))])
        # Schema: JSON-LD blocks
        schema_blocks = len(soup.find_all("script", {"type": "application/ld+json"}))
        # Images
        images = soup.find_all("img")
        img_with_alt = [i for i in images if i.get("alt")]
        # Title and meta
        title_tag = soup.find("title")
        title_text = title_tag.get_text() if title_tag else ""
        meta_desc_tag = soup.find("meta", {"name": "description"})
        meta_desc = meta_desc_tag.get("content", "") if meta_desc_tag else ""
        table_rows_all = [len(t.find_all("tr")) for t in tables]
    else:
        # Fallback: regex-based parsing (works on markdown-style plain text too)
        plain = re.sub(r'<[^>]+>', ' ', html)
        all_text = plain
        h1_tags  = re.findall(r'(?:H1:|<h1[^>]*>)(.*?)(?:\n|</h1>)', html, re.I)
        h2_tags  = re.findall(r'(?:H2:|<h2[^>]*>)(.*?)(?:\n|</h2>)', html, re.I)
        h3_tags  = re.findall(r'(?:H3:|<h3[^>]*>)(.*?)(?:\n|</h3>)', html, re.I)
        h4_tags  = re.findall(r'(?:H4:|<h4[^>]*>)(.*?)(?:\n|</h4>)', html, re.I)
        tables   = re.findall(r'<table[\s\S]*?</table>', html, re.I)
        table_rows_all = [len(re.findall(r'<tr', t, re.I)) for t in tables]
        lists_ul = re.findall(r'<ul[\s\S]*?</ul>', html, re.I)
        lists_ol = re.findall(r'<ol[\s\S]*?</ol>', html, re.I)
        anchors  = re.findall(r'<a[^>]*>(.*?)</a>', html, re.I)
        faq_count = len(re.findall(r'(?:FAQ|Q:|A:)', html, re.I))
        schema_blocks = len(re.findall(r'application/ld\+json', html, re.I))
        images = re.findall(r'<img[^>]*>', html, re.I)
        img_with_alt = [i for i in images if 'alt=' in i.lower()]
        title_text = (re.findall(r'<title>(.*?)</title>', html, re.I) or [""])[0]
        meta_desc_raw = re.findall(r'<meta[^>]*name=["\']description["\'][^>]*content=["\']([^"\']+)', html, re.I)
        meta_desc = meta_desc_raw[0] if meta_desc_raw else ""

    # ── Word count
    words_all = re.findall(r'\b\w+\b', plain)
    total_words = len(words_all)

    # ── Sentences
    sentences_all = [s.strip() for s in re.split(r'[.!?]+', plain) if s.strip() and len(s.split()) >= 3]
    total_sentences = len(sentences_all)
    long_sentences = [s for s in sentences_all if len(s.split()) > 30]
    avg_sent_len = round(sum(len(s.split()) for s in sentences_all) / max(1, total_sentences), 1)
    long_sent_pct = round(len(long_sentences) / max(1, total_sentences) * 100, 1)

    # ── Flesch
    flesch = _flesch_ease(plain[:5000])  # sample for speed

    # ── KW metrics
    if kw_lower:
        kw_count_total = len(re.findall(re.escape(kw_lower), plain.lower()))
        kw_density_pct = round(kw_count_total / max(1, total_words) * 100, 2)
        kw_in_h1 = sum(1 for h in h1_tags if kw_lower in h.lower())
        kw_in_h2 = sum(1 for h in h2_tags if kw_lower in h.lower())
        kw_in_h3 = sum(1 for h in h3_tags if kw_lower in h.lower())
        kw_in_title = 1 if kw_lower in title_text.lower() else 0
        first_100_words = " ".join(words_all[:100]).lower()
        kw_in_first100 = 1 if kw_lower in first_100_words else 0
        kw_in_anchors = sum(1 for a in anchors if kw_lower in a.lower())
        kw_in_h2_pct = round(kw_in_h2 / max(1, len(h2_tags)) * 100, 1)
    else:
        kw_count_total = kw_density_pct = kw_in_h1 = kw_in_h2 = kw_in_h3 = 0
        kw_in_title = kw_in_first100 = kw_in_anchors = kw_in_h2_pct = 0

    # ── LSI: unique non-stop words > 4 chars
    stop_words = {
        'this','that','with','have','from','they','been','were','will','would','could',
        'should','their','there','about','which','when','what','your','more','also',
        'some','into','than','then','each','only','most','over','such','after','before',
        'where','while','those','these','other','being','having','doing','through',
    }
    word_freq = {}
    for w in words_all:
        w_lower = w.lower()
        if len(w_lower) > 4 and w_lower not in stop_words and not w_lower.isdigit():
            word_freq[w_lower] = word_freq.get(w_lower, 0) + 1
    lsi_unique = len(word_freq)
    lsi_total  = sum(word_freq.values())

    # ── Table metrics
    table_count = len(tables)
    table_rows_avg = round(sum(table_rows_all) / max(1, table_count), 1) if table_rows_all else 0

    # ── List metrics
    list_count = len(lists_ul) + len(lists_ol)
    if _BS_OK:
        all_list_items = []
        for ul in soup.find_all(["ul", "ol"]):
            all_list_items.append(len(ul.find_all("li")))
        list_items_avg = round(sum(all_list_items) / max(1, len(all_list_items)), 1) if all_list_items else 0
    else:
        li_counts = [len(re.findall(r'<li', lst, re.I)) for lst in (lists_ul + lists_ol)]
        list_items_avg = round(sum(li_counts) / max(1, len(li_counts)), 1) if li_counts else 0

    return {
        "total_words":     total_words,
        "h1_count":        len(h1_tags),
        "h2_count":        len(h2_tags),
        "h3_count":        len(h3_tags),
        "h4_count":        len(h4_tags),
        "table_count":     table_count,
        "table_rows_avg":  table_rows_avg,
        "list_count":      list_count,
        "list_items_avg":  list_items_avg,
        "faq_count":       faq_count,
        "schema_blocks":   schema_blocks,
        "image_count":     len(images) if _BS_OK else len(re.findall(r'<img', html, re.I)),
        "images_with_alt": len(img_with_alt),
        "kw_count":        kw_count_total,
        "kw_density_pct":  kw_density_pct,
        "kw_in_h1":        kw_in_h1,
        "kw_in_h2":        kw_in_h2,
        "kw_in_h3":        kw_in_h3,
        "kw_in_h2_pct":    kw_in_h2_pct,
        "kw_in_title":     kw_in_title,
        "kw_in_first100":  kw_in_first100,
        "kw_in_anchors":   kw_in_anchors,
        "lsi_unique":      lsi_unique,
        "lsi_total":       lsi_total,
        "total_sentences": total_sentences,
        "avg_sent_len":    avg_sent_len,
        "long_sent_pct":   long_sent_pct,
        "flesch_ease":     flesch,
        "title_text":      title_text,
        "meta_desc":       meta_desc,
    }


async def get_targets_from_serp(keyword: str, geo: str, client) -> dict:
    """
    v3 CORE: Build dynamic content targets by:
      1. Fetching top-10 SERP URLs from Ahrefs API
      2. Scraping each URL HTML
      3. Extracting _extract_metrics_from_html() for each
      4. Computing MEDIAN of top-3 accessible sites per metric
      5. Applying Google-guideline hard caps

    Returns:
      {
        "word_count":     (min, max),
        "h2_count":       (min, max),
        "h3_count":       (min, max),
        "table_count":    (min, max),
        "list_count":     (min, max),
        "faq_count":      (min, max),
        "kw_density":     (min, max),
        "avg_sent_len":   (min, max),
        "long_sent_pct":  (min, max),
        "flesch_ease":    (min, max),
        "kw_in_h2_min":   N,
        "kw_in_h2_pct":   (min, max),
        "kw_in_anchors":  (min, max),
        "lsi_unique":     (min, max),
        "lsi_total":      (min, max),
        "schema_blocks":  (min, max),
        "table_rows_avg": (min, max),
        "list_items_avg": (min, max),
        "_serp_urls":     [...],
        "_metrics_raw":   [...],
      }
    Falls back to sane defaults if Ahrefs is unavailable.
    """
    geo_map = {
        "uk": "gb", "united kingdom": "gb", "great britain": "gb",
        "us": "us", "usa": "us", "united states": "us",
        "de": "de", "germany": "de", "deutschland": "de",
        "nl": "nl", "netherlands": "nl",
        "it": "it", "italy": "it",
        "pt": "pt", "portugal": "pt",
        "br": "br", "brazil": "br",
        "ca": "ca", "canada": "ca",
        "au": "au", "australia": "au",
    }
    geo_code = geo_map.get(geo.lower().strip(), geo.lower()[:2])

    import json as _json_cache, time as _time_cache
    _cache_key = re.sub(r"[^a-z0-9_]", "_", f"{keyword}_{geo_code}".lower())
    _cache_path = f"/tmp/ahrefs_cache_{_cache_key}.json"
    _CACHE_TTL = 86400  # 24 hours

    print(f"  🌐 v3 SERP: fetching Ahrefs top-10 for '{keyword}' / {geo_code.upper()}")

    serp_urls = []
    # ── Ahrefs cache (24h TTL) ───────────────────────────────────────────────
    _cache_hit = False
    if os.path.exists(_cache_path):
        try:
            _cache_data = _json_cache.loads(open(_cache_path).read())
            if _time_cache.time() - _cache_data.get("ts", 0) < _CACHE_TTL:
                serp_urls = _cache_data.get("urls", [])
                print(f"  ✅ Ahrefs cache HIT ({len(serp_urls)} URLs, age {int((_time_cache.time()-_cache_data['ts'])/60)}m)")
                _cache_hit = True
        except Exception:
            pass
    if not _cache_hit:
        try:
            ahrefs_url = (
                f"https://api.ahrefs.com/v3/serp-overview/serp-overview"
                f"?select=url,position&keyword={keyword}&country={geo_code}"
            )
            r = await client.get(
                ahrefs_url,
                headers={"Authorization": f"Bearer {AHREFS_TOKEN}", "Accept": "application/json"},
                timeout=15.0,
            )
            if r.status_code == 200:
                data = r.json()
                positions = data.get("serp_overview", data.get("positions", []))
                serp_urls = [p["url"] for p in positions if p.get("url")][:10]
                print(f"  ✅ Ahrefs returned {len(serp_urls)} SERP URLs")
                # Save cache
                try:
                    open(_cache_path, "w").write(_json_cache.dumps({"ts": _time_cache.time(), "urls": serp_urls}))
                    print(f"  💾 Ahrefs cache saved → {_cache_path}")
                except Exception:
                    pass
            else:
                print(f"  ⚠️ Ahrefs HTTP {r.status_code}: {r.text[:120]}")
        except Exception as e:
            print(f"  ⚠️ Ahrefs API error: {e}")

    # Fallback: use task competitors if Ahrefs returned nothing
    if not serp_urls:
        print("  ↩️  Ahrefs empty — will use task.competitors as SERP proxy")

    # ── Scrape each URL ──────────────────────────────────────────────────────
    # ── Parallel scrape top-10 with 5s timeout per site ────────────────────
    import asyncio as _aio

    async def _scrape_one(url: str) -> Optional[dict]:
        try:
            from smart_fetcher import fetch_html as _smart_fetch
            _html, _method = await _smart_fetch(url, timeout=10, use_playwright=False)
            if _html:
                metrics = _extract_metrics_from_html(_html, keyword)
                if metrics:
                    metrics["_html"] = _html  # store raw HTML for TF-IDF / heading analysis
                return metrics
        except Exception:
            pass
        return None

    t_scrape_start = _time_cache.time()
    scrape_tasks = [_scrape_one(u) for u in serp_urls[:10]]
    raw_list = await _aio.gather(*scrape_tasks, return_exceptions=True)
    metrics_raw = [m for m in raw_list if isinstance(m, dict) and m and m.get("total_words", 0) > 200]
    print(f"  📊 Scraped {len(metrics_raw)}/{len(serp_urls)} SERP pages in {_time_cache.time()-t_scrape_start:.1f}s")

    if not metrics_raw:
        print("  ⚠️ No SERP data — using v3 default targets")
        return _default_targets()

    # ── Compute MEDIAN over top-3 (sorted by word count desc) ───────────────
    metrics_raw.sort(key=lambda m: m["total_words"], reverse=True)
    top3 = metrics_raw[:3]

    def _median(vals):
        v = sorted(v for v in vals if v is not None)
        if not v:
            return 0
        n = len(v)
        return v[n // 2] if n % 2 == 1 else (v[n // 2 - 1] + v[n // 2]) / 2

    def _range(key, expand=0.15):
        """Median ± expand% as (min, max) range, rounded."""
        med = _median([m.get(key, 0) for m in top3])
        lo  = max(0, _math_v3.floor(med * (1 - expand)))
        hi  = _math_v3.ceil(med * (1 + expand))
        return (lo, hi)

    wc_med = _median([m["total_words"] for m in top3])

    # Build raw targets from medians
    targets = {
        "word_count":     (_math_v3.floor(wc_med * 0.9), _math_v3.ceil(wc_med * 1.1)),
        "h2_count":       _range("h2_count"),
        "h3_count":       _range("h3_count"),
        "table_count":    _range("table_count"),
        "table_rows_avg": _range("table_rows_avg"),
        "list_count":     _range("list_count"),
        "list_items_avg": _range("list_items_avg"),
        "faq_count":      _range("faq_count", expand=0.25),
        "schema_blocks":  _range("schema_blocks"),
        "total_sentences":_range("total_sentences"),
        "avg_sent_len":   _range("avg_sent_len", expand=0.2),
        "long_sent_pct":  (0.0, min(_LONG_SENT_MAX, _median([m["long_sent_pct"] for m in top3]) + 1.0)),
        "flesch_ease":    (max(_FLESCH_MIN, _range("flesch_ease")[0]), _range("flesch_ease")[1]),
        "lsi_unique":     _range("lsi_unique"),
        "lsi_total":      _range("lsi_total"),
        "kw_density":     (
            max(0.3, _median([m["kw_density_pct"] for m in top3]) * 0.8),
            min(_KW_DENSITY_MAX, _median([m["kw_density_pct"] for m in top3]) * 1.2),
        ),
        "kw_in_h2_min":   max(2, min(int(_median([m["kw_in_h2_pct"] for m in top3]) / 100 * max(8, int(_median([m.get("h2_count", 10) for m in top3])))), int(max(8, int(_median([m.get("h2_count", 10) for m in top3]))) * 0.25))),  # cap at 25% of H2 count
        "kw_in_h2_pct":   _range("kw_in_h2_pct"),
        "kw_in_anchors":  _range("kw_in_anchors"),
        # Binary targets
        "kw_in_h1":       1,
        "kw_in_title":    1,
        "kw_in_first100": 1,
        "h1_count":       1,
        # Meta
        "_serp_urls":     serp_urls,
        "_metrics_raw":   metrics_raw,
        "_serp_htmls":    [m.get("_html", "") for m in metrics_raw if m.get("_html")],
        "_n_competitors": len(top3),
    }

    print(f"  ✅ v3 targets computed (median of top-{len(top3)}): "
          f"words={targets['word_count']}, H2={targets['h2_count']}, "
          f"KW density={targets['kw_density'][0]:.1f}-{targets['kw_density'][1]:.1f}%")
    return targets


def _default_targets() -> dict:
    """Fallback static targets when Ahrefs/scraping fails."""
    return {
        "word_count":     (3500, 4500),
        "h2_count":       (14, 18),
        "h3_count":       (25, 35),
        "table_count":    (8, 14),
        "table_rows_avg": (6, 9),
        "list_count":     (6, 10),
        "list_items_avg": (4, 7),
        "faq_count":      (14, 20),
        "schema_blocks":  (3, 6),
        "total_sentences":(200, 350),
        "avg_sent_len":   (10, 18),
        "long_sent_pct":  (0.0, 5.0),
        "flesch_ease":    (10.0, 40.0),
        "lsi_unique":     (300, 600),
        "lsi_total":      (400, 1000),
        "kw_density":     (0.5, 3.0),
        "kw_in_h2_min":   2,  # 20-25% of avg 8-10 H2 = 2 headings minimum
        "kw_in_h2_pct":   (20, 30),
        "kw_in_anchors":  (3, 8),
        "kw_in_h1":       1,
        "kw_in_title":    1,
        "kw_in_first100": 1,
        "h1_count":       1,
        "_serp_urls":     [],
        "_metrics_raw":   [],
    }


def _in_range(value, target) -> bool:
    """Check if value is within target (tuple range or scalar)."""
    if isinstance(target, tuple):
        return target[0] <= value <= target[1]
    return value >= target


def _range_str(target) -> str:
    if isinstance(target, tuple):
        lo, hi = target
        if isinstance(lo, float) or isinstance(hi, float):
            return f"{lo:.1f}–{hi:.1f}"
        return f"{lo}–{hi}"
    return str(target)


def analyze_content(html: str, task: dict, targets: dict) -> dict:
    """
    Full post-generation content audit.

    Returns:
      {
        "metrics":  {key: value, ...},
        "report":   "📊 ПРОВЕРКА ТЕКСТА:\n✅ ...\n❌ ...",
        "issues":   [{key, value, target, fixable}, ...],
        "warnings": [{key, value, target}, ...],
        "ok":       bool,
      }
    """
    keyword = task.get("keywords", [""])[0] if task.get("keywords") else ""
    metrics = _extract_metrics_from_html(html, keyword)

    # ── Define checks ─────────────────────────────────────────────────────────
    checks = [
        # (metric_key, label, target_key, fixable)
        ("total_words",      "Words",               "word_count",      False),
        ("h1_count",         "H1",                  "h1_count",        True),
        ("h2_count",         "H2",                  "h2_count",        False),
        ("h3_count",         "H3",                  "h3_count",        False),
        ("table_count",      "Tables",              "table_count",     False),
        ("table_rows_avg",   "Table rows avg",      "table_rows_avg",  False),
        ("list_count",       "Lists",               "list_count",      False),
        ("faq_count",        "FAQ items",           "faq_count",       False),
        ("schema_blocks",    "Schema blocks",       "schema_blocks",   False),
        ("kw_density_pct",   "KW density %",        "kw_density",      False),
        ("kw_in_h1",         "KW in H1",            "kw_in_h1",        True),
        ("kw_in_h2",         "KW in H2",            "kw_in_h2_min",    True),
        ("kw_in_h2_pct",     "KW in H2 %",          "kw_in_h2_pct",    False),
        ("kw_in_title",      "KW in Title",         "kw_in_title",     False),
        ("kw_in_first100",   "KW in first 100w",    "kw_in_first100",  True),
        ("kw_in_anchors",    "KW in anchors",       "kw_in_anchors",   False),
        ("lsi_unique",       "LSI unique terms",    "lsi_unique",      False),
        ("lsi_total",        "LSI total uses",      "lsi_total",       False),
        ("avg_sent_len",     "Avg sentence len",    "avg_sent_len",    False),
        ("long_sent_pct",    "Long sentences %",    "long_sent_pct",   False),
        ("flesch_ease",      "Flesch Reading Ease", "flesch_ease",     False),
        ("total_sentences",  "Sentences total",     "total_sentences", False),
    ]

    report_lines = ["📊 ПРОВЕРКА ТЕКСТА (v3 AHREFS-DRIVEN):"]
    issues   = []
    warnings = []

    for metric_key, label, target_key, fixable in checks:
        val    = metrics.get(metric_key, 0)
        target = targets.get(target_key)
        if target is None:
            continue

        ok = _in_range(val, target)
        tstr = _range_str(target)

        if ok:
            report_lines.append(f"  ✅ {label}: {val} (цель: {tstr})")
        else:
            # Distinguish hard fail vs warning
            # Hard fail: KW in H1, KW in first 100, H1 count
            is_critical = metric_key in ("kw_in_h1", "kw_in_first100", "h1_count", "kw_in_h2")
            if is_critical or fixable:
                report_lines.append(f"  ❌ {label}: {val} (цель: {tstr})" + (" → ИСПРАВЛЕНО" if fixable else ""))
                issues.append({"key": metric_key, "value": val, "target": target, "fixable": fixable, "label": label})
            else:
                report_lines.append(f"  ⚠️  {label}: {val} (цель: {tstr})")
                warnings.append({"key": metric_key, "value": val, "target": target, "label": label})


    # Append search intent / snippet info if available in targets
    intent = targets.get("_search_intent")
    if intent:
        report_lines.append(f"  \U0001f3af Search Intent: {intent.upper()}")

    # Append TF-IDF coverage note
    tfidf_terms = targets.get("_tfidf_terms", [])
    if tfidf_terms:
        top_terms = [t for t, _ in tfidf_terms[:10]]
        html_lower = html.lower()
        covered = sum(1 for t in top_terms if t.lower() in html_lower)
        report_lines.append(f"  \U0001f511 TF-IDF top-10 coverage: {covered}/10 terms present in text")

    report = "\n".join(report_lines)
    ok = len([i for i in issues if not i["fixable"]]) == 0

    return {
        "metrics":  metrics,
        "report":   report,
        "issues":   issues,
        "warnings": warnings,
        "ok":       ok,
    }


def fix_content_issues(html: str, issues: list, task: dict, keyword: str = "") -> str:
    """
    Auto-fix critical content issues where possible.

    Fixable issues:
      - kw_in_h1:       add keyword to H1 if missing
      - kw_in_first100: insert keyword into first paragraph
      - kw_in_h2:       add keyword to N earliest H2s to hit target
    """
    if not issues or not html:
        return html

    result = html
    kw = keyword or (task.get("keywords", [""])[0] if task.get("keywords") else "")
    kw_lower = kw.lower().strip()

    for issue in issues:
        if not issue.get("fixable") or not kw_lower:
            continue
        key = issue["key"]
        target = issue["target"]

        # ── Fix: KW in H1 ──────────────────────────────────────────────────
        if key == "kw_in_h1":
            # Pattern: "H1: some text" OR "<h1>some text</h1>"
            h1_match = re.search(r'(H1:\s*)([^\n]+)', result) or re.search(r'(<h1[^>]*>)(.*?)(</h1>)', result, re.I | re.S)
            if h1_match:
                if kw_lower not in h1_match.group(0).lower():
                    if '<h1' in h1_match.group(0).lower():
                        old = h1_match.group(0)
                        inner = h1_match.group(2).rstrip()
                        new_inner = f"{inner} — {kw}"
                        result = result.replace(old, f"{h1_match.group(1)}{new_inner}{h1_match.group(3)}", 1)
                    else:
                        old = h1_match.group(0)
                        result = result.replace(old, f"{h1_match.group(1)}{h1_match.group(2).rstrip()} — {kw}", 1)

        # ── Fix: KW in first 100 words ────────────────────────────────────
        elif key == "kw_in_first100":
            # Find first paragraph-like content line and prepend KW mention
            para_match = re.search(r'(<p[^>]*>)(.*?)(</p>)', result, re.I | re.S)
            if para_match:
                old_para = para_match.group(0)
                inner = para_match.group(2).strip()
                if kw_lower not in inner.lower():
                    new_inner = f"{kw.capitalize()} — {inner}"
                    result = result.replace(old_para, f"{para_match.group(1)}{new_inner}{para_match.group(3)}", 1)
            else:
                # Fallback: markdown-style plain text — first non-heading line
                lines = result.split('\n')
                for i, line in enumerate(lines):
                    stripped = line.strip()
                    if stripped and not re.match(r'^H[1-4]:|^<h[1-4]|^Meta |^#', stripped) and len(stripped) > 30:
                        if kw_lower not in stripped.lower():
                            lines[i] = f"{kw.capitalize()} — {stripped}"
                            break
                result = '\n'.join(lines)

        # ── Fix: KW in H2 ─────────────────────────────────────────────────
        elif key == "kw_in_h2":
            target_count = target if isinstance(target, int) else int(target)
            current_kw_h2 = len(re.findall(re.escape(kw_lower), '\n'.join(re.findall(r'H2:[^\n]+|<h2[^>]*>.*?</h2>', result, re.I)), re.I))
            needed = max(0, target_count - current_kw_h2)
            # Find H2s that DON'T have the keyword and add it
            def _add_kw_to_h2(m):
                nonlocal needed
                if needed > 0 and kw_lower not in m.group(0).lower():
                    needed -= 1
                    return m.group(0).rstrip() + f" — {kw}"
                return m.group(0)
            result = re.sub(r'H2:[^\n]+', _add_kw_to_h2, result)
            result = re.sub(r'<h2[^>]*>.*?</h2>', _add_kw_to_h2, result, flags=re.I | re.S)

    return result


def build_serp_targets_prompt_block(targets: dict, keyword: str) -> str:
    """
    Build the prompt injection block from SERP-derived targets.
    This replaces build_keyword_targets_block() for v3.
    """
    if not targets or not targets.get("word_count"):
        return ""

    wc = targets["word_count"]
    h2 = targets["h2_count"]
    h3 = targets["h3_count"]
    tb = targets["table_count"]
    ls = targets["list_count"]
    fq = targets["faq_count"]
    kd = targets["kw_density"]
    as_ = targets["avg_sent_len"]
    lp  = targets["long_sent_pct"]
    fl  = targets["flesch_ease"]
    lsi = targets["lsi_unique"]
    sc  = targets["schema_blocks"]
    kh2 = targets.get("kw_in_h2_min", 3)
    ka  = targets.get("kw_in_anchors", (3, 8))

    def r(t):
        return _range_str(t)

    return f"""
═══ 🎯 v3 AHREFS-DRIVEN TARGETS (ОБЯЗАТЕЛЬНО — на основе медианы топ-3 SERP) ═══
⚠️ Все параметры вычислены из реального анализа конкурентов в топ Google.
Отклонение от диапазонов = текст не пройдёт проверку QA.

▸ СТРУКТУРА КОНТЕНТА:
  - Total words:      {r(wc)} слов (ОБЯЗАТЕЛЬНО)
  - H1 count:         ровно 1 (ОБЯЗАТЕЛЬНО)
  - H2 count:         {r(h2)} заголовков H2
  - H3 count:         {r(h3)} заголовков H3
  - Tables:           {r(tb)} таблиц (avg {r(targets.get('table_rows_avg',(6,9)))} строк каждая)
  - Lists (ul/ol):    {r(ls)} списков (avg {r(targets.get('list_items_avg',(4,7)))} пунктов)
  - FAQ items:        {r(fq)} вопросов в FAQ секции
  - Total sentences:  {r(targets.get('total_sentences',(200,350)))}
  - Schema blocks:    {r(sc)} JSON-LD блоков (WebPage, BreadcrumbList, FAQPage, Person)

▸ КЛЮЧЕВЫЕ СЛОВА «{keyword}»:
  - KW ОБЯЗАТЕЛЬНО в H1 (минимум 1 раз)
  - KW ОБЯЗАТЕЛЬНО в Title
  - KW ОБЯЗАТЕЛЬНО в первых 100 словах
  - KW в H2: 20-25% заголовков H2 содержат «{keyword}» (не более чем каждый 4-5 заголовок). Остальные H2 должны использовать синонимы и LSI термины
  - KW в якорях (anchor text): {r(ka)} раз
  - KW density: {r(kd)}% от общего числа слов (НЕ ВЫШЕ 3% — Google anti-spam cap)

▸ ЧИТАЕМОСТЬ:
  - Avg sentence length: {r(as_)} слов (НЕ короче 8, НЕ длиннее 21 в среднем)
  - Long sentences >30 слов: НЕ БОЛЕЕ {lp[1]:.1f}% от всех предложений
  - Flesch Reading Ease: {r(fl)} (минимум 10 — Google guideline)

▸ LSI / СЕМАНТИКА:
  - LSI unique terms:  {r(lsi)} уникальных значимых слов
  - LSI total uses:    {r(targets.get('lsi_total',(400,1000)))} суммарных употреблений

⚠️ ПОСЛЕ ГЕНЕРАЦИИ СКРИПТ АВТОМАТИЧЕСКИ ПРОВЕРИТ ВСЕ ПАРАМЕТРЫ.
❌-проблемы будут исправлены принудительно. ⚠️-отклонения попадут в отчёт.

"""

# ══════════════════════════════════════════════════════════════════════════════
# ▲▲▲  END v3 ADDITIONS  ▲▲▲

# ══════════════════════════════════════════════════════════════════════════════
# ▼▼▼  v3 EXTENDED MODULES  ▼▼▼
# ══════════════════════════════════════════════════════════════════════════════

# ─── MODULE 1: TF-IDF анализ конкурентов ─────────────────────────────────────

def _strip_html(html: str) -> str:
    """Simple HTML tag stripper."""
    import re as _re
    text = _re.sub(r'<script[^>]*>.*?</script>', ' ', html, flags=_re.DOTALL | _re.I)
    text = _re.sub(r'<style[^>]*>.*?</style>', ' ', text, flags=_re.DOTALL | _re.I)
    text = _re.sub(r'<[^>]+>', ' ', text)
    text = _re.sub(r'&[a-z]+;', ' ', text)
    text = _re.sub(r'\s+', ' ', text)
    return text.strip()


def extract_tfidf_terms(competitor_htmls: list, top_n: int = 50, lang: str = "en") -> list:
    """
    Извлекает топ-N значимых терминов из HTML конкурентов через TF-IDF.
    Возвращает [(term, score), ...] отсортированных по важности.
    Использует sklearn TfidfVectorizer или реализацию вручную если sklearn недоступен.
    Фильтрует стоп-слова по языку задачи.
    """
    import re as _re
    import math as _math

    if not competitor_htmls:
        return []

    # Strip HTML → plain text documents
    docs = [_strip_html(h) for h in competitor_htmls if h]
    docs = [d for d in docs if len(d) > 50]
    if not docs:
        return []

    # Common stop-word sets per language
    _STOP_EN = set("a an the and or but in on at to of for is are was were be been being have has had do does did will would could should may might must shall can need this that these those with from by about into through during before after above below between each few more most other some than then there they them their what which who when where why how all both each few more most other some such no nor not only own same so than too very just now".split())
    _STOP_NL = set("de het een en van in is op dat te zijn voor met als ook niet maar hij zij we zijn ze had kunnen worden door worden met zijn zijn dat".split())
    _STOP_IT = set("il la le lo gli i un una uno e di a in è che si per con non da come sono ha al del della dei alle degli".split())
    _STOP_DE = set("der die das den dem des ein eine einer einem einen eines und ist in zu von mit für an auf aus bei durch nach über unter vor zwischen".split())
    _STOP_PT = set("o a os as um uma de em por para com que não se na no ao da do dos das".split())

    _stop_map = {"nl": _STOP_NL, "it": _STOP_IT, "de": _STOP_DE, "pt": _STOP_PT}
    stop_words = _STOP_EN | _stop_map.get(lang.lower()[:2], set())

    try:
        from sklearn.feature_extraction.text import TfidfVectorizer
        vec = TfidfVectorizer(
            ngram_range=(1, 2),
            max_features=500,
            stop_words=list(stop_words),
            min_df=1,
            sublinear_tf=True,
        )
        tfidf_matrix = vec.fit_transform(docs)
        feature_names = vec.get_feature_names_out()
        # Average TF-IDF score across all documents
        scores = tfidf_matrix.mean(axis=0).A1
        ranked = sorted(zip(feature_names, scores), key=lambda x: x[1], reverse=True)
        # Filter short terms and pure numbers
        ranked = [(t, s) for t, s in ranked if len(t) > 2 and not t.replace(' ', '').isdigit()]
        return ranked[:top_n]
    except ImportError:
        pass

    # Manual TF-IDF fallback
    def _tokenize(text):
        tokens = _re.findall(r'\b[a-zA-ZÀ-ÿ]{3,}\b', text.lower())
        return [t for t in tokens if t not in stop_words]

    tokenized_docs = [_tokenize(d) for d in docs]
    N = len(tokenized_docs)

    # Build DF (document frequency)
    df = {}
    for tokens in tokenized_docs:
        for term in set(tokens):
            df[term] = df.get(term, 0) + 1

    # Build TF per doc, then average TF-IDF
    term_scores = {}
    for tokens in tokenized_docs:
        if not tokens:
            continue
        tf = {}
        for t in tokens:
            tf[t] = tf.get(t, 0) + 1
        total = len(tokens)
        for term, count in tf.items():
            tf_val = count / total
            idf_val = _math.log((N + 1) / (df.get(term, 0) + 1)) + 1
            term_scores[term] = term_scores.get(term, 0.0) + tf_val * idf_val

    ranked = sorted(term_scores.items(), key=lambda x: x[1], reverse=True)
    ranked = [(t, round(s / N, 5)) for t, s in ranked if len(t) > 2]
    return ranked[:top_n]


def build_tfidf_prompt_block(tfidf_terms: list) -> str:
    """Форматирует TF-IDF термины для вставки в промпт."""
    if not tfidf_terms:
        return ""
    high = [t for t, _ in tfidf_terms[:10]]
    medium = [t for t, _ in tfidf_terms[10:30]]
    low = [t for t, _ in tfidf_terms[30:50]]
    lines = ["\n═══ 🔑 LSI/SEMANTIC TERMS (использовать в тексте) ═══"]
    if high:
        lines.append(f"  High priority (обязательно): {', '.join(high)}")
    if medium:
        lines.append(f"  Medium priority (желательно): {', '.join(medium)}")
    if low:
        lines.append(f"  Low priority (по возможности): {', '.join(low)}")
    lines.append("  ⚠️ Используй естественно — не keyword stuffing!\n")
    return "\n".join(lines)


# ─── MODULE 2: Анализ структуры H2/H3 конкурентов ────────────────────────────

def extract_competitor_headings(html: str) -> dict:
    """
    Извлекает все H1/H2/H3 заголовки конкурента.
    Возвращает {"h1": [...], "h2": [...], "h3": [...]}
    """
    import re as _re
    result = {"h1": [], "h2": [], "h3": []}
    for level in (1, 2, 3):
        pattern = rf'<h{level}[^>]*>(.*?)</h{level}>'
        matches = _re.findall(pattern, html, _re.DOTALL | _re.I)
        result[f"h{level}"] = [
            _re.sub(r'<[^>]+>', '', m).strip()[:120]
            for m in matches
            if m.strip()
        ]
    return result


def build_topic_map(competitor_headings_list: list) -> dict:
    """
    Из заголовков топ-3 конкурентов строит тематическую карту.
    Находит общие темы H2 (которые есть у 2+ конкурентов).
    Возвращает {"must_cover": [...], "optional": [...]}
    """
    import re as _re

    if not competitor_headings_list:
        return {"must_cover": [], "optional": []}

    n = len(competitor_headings_list)

    def _normalize(heading: str) -> str:
        return _re.sub(r'[^\w\s]', '', heading.lower()).strip()

    def _similarity(a: str, b: str) -> float:
        """Simple word overlap similarity."""
        wa = set(a.split())
        wb = set(b.split())
        if not wa or not wb:
            return 0.0
        return len(wa & wb) / max(len(wa), len(wb))

    # Collect all H2s from all competitors
    all_h2s = []
    for i, headings in enumerate(competitor_headings_list):
        for h2 in headings.get("h2", []):
            all_h2s.append((i, _normalize(h2), h2))

    # Group similar headings
    groups = []  # [(canonical_text, [competitor_indices], [original_headings])]
    used = set()

    for idx, (comp_i, norm, orig) in enumerate(all_h2s):
        if idx in used:
            continue
        group_comps = {comp_i}
        group_origs = [orig]
        used.add(idx)

        for idx2, (comp_j, norm2, orig2) in enumerate(all_h2s):
            if idx2 in used or comp_j in group_comps:
                continue
            if _similarity(norm, norm2) >= 0.4:
                group_comps.add(comp_j)
                group_origs.append(orig2)
                used.add(idx2)

        groups.append((orig, group_comps, group_origs))

    # Sort by frequency
    groups.sort(key=lambda g: len(g[1]), reverse=True)

    must_cover = []
    optional = []
    for canonical, comps, origs in groups:
        freq = len(comps)
        entry = {
            "topic": canonical,
            "variants": origs[:3],
            "freq": freq,
            "total": n,
        }
        if freq >= 2:
            must_cover.append(entry)
        else:
            optional.append(entry)

    return {
        "must_cover": must_cover[:12],
        "optional": optional[:8],
    }


def build_topic_map_prompt_block(topic_map: dict) -> str:
    """Форматирует тематическую карту для вставки в промпт."""
    if not topic_map or not topic_map.get("must_cover"):
        return ""
    lines = ["\n═══ 🗺️ TOPIC MAP (обязательно покрыть эти темы в H2) ═══"]
    n_total = topic_map["must_cover"][0]["total"] if topic_map["must_cover"] else 3
    for entry in topic_map["must_cover"]:
        freq = entry["freq"]
        topic = entry["topic"]
        lines.append(f"  ✅ {topic} — у {freq}/{n_total} конкурентов")
    if topic_map.get("optional"):
        lines.append("  Дополнительные темы (по возможности):")
        for entry in topic_map["optional"][:5]:
            lines.append(f"  ➕ {entry['topic']}")
    lines.append("  ⚠️ ЗАПРЕЩЕНО копировать заголовки дословно — перефразируй!\n")
    return "\n".join(lines)


# ─── MODULE 3: Featured Snippet оптимизация ──────────────────────────────────

_INTENT_PATTERNS = {
    "how-to": [
        r'\b(how to|how do|how can|пошагово|как|step[- ]by[- ]step|guide|tutorial|инструкция)\b',
        r'\b(hoe|wie man|come fare|como)\b',
    ],
    "comparison": [
        r'\b(vs\.?|versus|compare|best|top|лучш|сравн|рейтинг|ranking|rated|compared)\b',
        r'\b(vergelijk|vergleich|migliori|melhores|mejores)\b',
    ],
    "review": [
        r'\b(review|отзыв|обзор|experience|opinion|rating|erfahrung|recensione|avaliação|reseña)\b',
        r'\b(legit|safe|trusted|betrouwbaar|sicuro|confiável)\b',
    ],
    "informational": [
        r'\b(what is|what are|why|when|who|where|что такое|что это|зачем|почему|когда)\b',
        r'\b(wat is|was ist|che cos|o que é|qué es)\b',
    ],
}


def detect_search_intent(keyword: str, geo: str = "us") -> str:
    """
    Определяет intent запроса по ключевому слову.
    Returns: "informational" | "comparison" | "review" | "how-to"
    """
    import re as _re
    kw_lower = keyword.lower().strip()

    for intent, patterns in _INTENT_PATTERNS.items():
        for pattern in patterns:
            if _re.search(pattern, kw_lower, _re.I):
                return intent

    # Fallback heuristics
    words = kw_lower.split()
    if len(words) >= 4:
        return "informational"
    if any(w in kw_lower for w in ["casino", "online", "bonus", "slot"]):
        return "comparison"
    return "informational"


def build_featured_snippet_block(keyword: str, intent: str, lang: str = "en") -> str:
    """
    Генерирует инструкцию для Featured Snippet блока.
    Returns: строка инструкций для вставки в промпт.
    """
    lang = lang.lower()[:2]

    if intent == "how-to":
        snippet_instructions = (
            f"Numbered list (5-7 steps), each step 10-20 words. "
            f"Start with action verb. First item answers 'how to {keyword}' directly."
        )
        html_hint = (
            f"<ol><li>Step one...</li><li>Step two...</li>...</ol>"
        )
    elif intent == "comparison":
        snippet_instructions = (
            f"Comparison table (3-5 columns: Name/Rating/Bonus/License/Feature). "
            f"2-3 sentence intro before table answering 'which is best'."
        )
        html_hint = (
            f"<table><thead><tr><th>Casino</th><th>Rating</th><th>Bonus</th></tr></thead>"
            f"<tbody>...</tbody></table>"
        )
    elif intent == "review":
        snippet_instructions = (
            f"2-3 sentence summary with rating, key pros (3 items), key cons (2 items). "
            f"Include overall score X/10. Start with '{keyword} is...' or '{keyword} offers...'."
        )
        html_hint = (
            f"<div class='snippet-review'><p>Rating: X/10</p>"
            f"<ul><li>Pro 1</li><li>Pro 2</li></ul></div>"
        )
    else:  # informational
        snippet_instructions = (
            f"Definition paragraph: 40-60 words, plain language, directly answers 'what is {keyword}'. "
            f"No jargon. First sentence = direct answer. Second/third = context."
        )
        html_hint = (
            f"<div class='snippet-definition'><p>Direct answer here in 40-60 words.</p></div>"
        )

    return (
        f"\n═══ ⭐ FEATURED SNIPPET OPTIMIZATION ═══\n"
        f"  Intent detected: {intent.upper()}\n"
        f"  📌 Добавь СРАЗУ ПОСЛЕ H1 (перед первым параграфом) Featured Snippet блок:\n"
        f"  Format: {snippet_instructions}\n"
        f"  HTML structure hint: {html_hint}\n"
        f"  Цель: попасть в Position Zero Google. Блок должен быть самодостаточным (ответ без клика).\n"
    )


# ─── INTEGRATION HELPER: расширенный build_serp_targets_prompt_block ─────────

def build_extended_serp_prompt_block(
    targets: dict,
    keyword: str,
    competitor_htmls: list = None,
    lang: str = "en",
    geo: str = "us",
) -> str:
    """
    Расширенный промпт-блок: SERP targets + TF-IDF + Topic Map + Featured Snippet.
    Вызывается вместо build_serp_targets_prompt_block() когда есть HTML конкурентов.
    """
    base_block = build_serp_targets_prompt_block(targets, keyword)

    extra_blocks = []

    # TF-IDF block
    if competitor_htmls:
        try:
            tfidf_terms = extract_tfidf_terms(competitor_htmls, top_n=50, lang=lang)
            if tfidf_terms:
                extra_blocks.append(build_tfidf_prompt_block(tfidf_terms))
                # Store in targets for downstream use
                targets["_tfidf_terms"] = tfidf_terms
        except Exception as _e:
            print(f"  ⚠️ TF-IDF error: {_e}")

        # Topic Map block
        try:
            headings_list = [extract_competitor_headings(h) for h in competitor_htmls[:5] if h]
            topic_map = build_topic_map(headings_list)
            if topic_map.get("must_cover"):
                extra_blocks.append(build_topic_map_prompt_block(topic_map))
                targets["_topic_map"] = topic_map
        except Exception as _e:
            print(f"  ⚠️ Topic map error: {_e}")

    # Featured Snippet block
    try:
        intent = detect_search_intent(keyword, geo)
        snippet_block = build_featured_snippet_block(keyword, intent, lang)
        extra_blocks.append(snippet_block)
        targets["_search_intent"] = intent
    except Exception as _e:
        print(f"  ⚠️ Featured snippet error: {_e}")

    return base_block + "\n".join(extra_blocks)

# ══════════════════════════════════════════════════════════════════════════════
# ▲▲▲  END v3 EXTENDED MODULES  ▲▲▲
# ══════════════════════════════════════════════════════════════════════════════
# ══════════════════════════════════════════════════════════════════════════════
async def analyze_competitor_structure(client, urls, keyword=""):
    """
    Deep structural analysis of competitor pages:
    - H1/H2/H3 heading structure
    - Content block types (lists, tables, callouts)
    - First-person experience phrases
    - E-E-A-T signals
    - Avg paragraph length
    Returns structural recommendations for prompt injection.
    """
    import re as _re
    
    if not urls:
        return None
    
    all_h2s = []
    all_h3s = []
    has_table = 0
    has_lists = 0
    has_callouts = 0
    experience_phrases = []
    total_sites = 0
    
    # E-E-A-T experience phrases patterns
    experience_patterns = [
        r"\bwe tested\b", r"\bwe found\b", r"\bwe noticed\b", r"\bduring our test\b",
        r"\bin our experience\b", r"\bwe personally\b", r"\bour review\b", r"\bwe verified\b",
        r"\bwij hebben\b", r"\bons team\b", r"\bwij testen\b", r"\bwe hebben\b",
        r"\bnashim\b", r"\bmы протестировали\b", r"\bнаши эксперты\b",
        r"\bnoi abbiamo\b", r"\bwir haben\b", r"\bnos expertos\b",
        r"\bduring testing\b", r"\bpersonally tested\b", r"\breal money\b",
    ]
    
    async def fetch_structure(url):
        for attempt in range(1):
            try:
                from smart_fetcher import fetch_html as _smart_fetch
                html, _method = await _smart_fetch(url, timeout=12, use_playwright=False)
                if html:
                    # Extract headings
                    h1s = [_re.sub(r"<[^>]+>", "", h).strip()[:80] for h in _re.findall(r"<h1[^>]*>(.*?)</h1>", html, _re.DOTALL|_re.I)]
                    h2s = [_re.sub(r"<[^>]+>", "", h).strip()[:80] for h in _re.findall(r"<h2[^>]*>(.*?)</h2>", html, _re.DOTALL|_re.I)]
                    h3s = [_re.sub(r"<[^>]+>", "", h).strip()[:60] for h in _re.findall(r"<h3[^>]*>(.*?)</h3>", html, _re.DOTALL|_re.I)]
                    # Block types
                    has_tbl = bool(_re.search(r"<table", html, _re.I))
                    has_lst = bool(_re.search(r"<(?:ul|ol)\s", html, _re.I))
                    has_cll = bool(_re.search(r'class="[^"]*(?:callout|tip|warning|note|highlight|info-box|fs-paragraph)[^"]*"', html, _re.I))
                    # E-E-A-T phrases
                    text_lower = html.lower()
                    found_exp = [p for p in experience_patterns if _re.search(p, text_lower)]
                    return {
                        "h1": h1s[:1], "h2": h2s[:8], "h3": h3s[:6],
                        "table": has_tbl, "list": has_lst, "callout": has_cll,
                        "experience": found_exp,
                        "h2_count": len(h2s), "h3_count": len(h3s),
                    }
            except Exception:
                pass
        return None
    
    import asyncio
    tasks = [fetch_structure(url) for url in urls[:8]]
    raw = await asyncio.gather(*tasks)
    results = [r for r in raw if r]
    
    if not results:
        return None
    
    n = len(results)
    # Aggregate
    for r in results:
        all_h2s.extend([h for h in r["h2"] if h and h.strip()])
        all_h3s.extend([h for h in r["h3"] if h and h.strip()])
        has_table += int(r["table"])
        has_lists += int(r["list"])
        has_callouts += int(r["callout"])
        experience_phrases.extend(r["experience"])
    
    avg_h2 = sum(r["h2_count"] for r in results) / n
    avg_h3 = sum(r["h3_count"] for r in results) / n
    table_pct = has_table / n * 100
    list_pct = has_lists / n * 100
    callout_pct = has_callouts / n * 100
    
    # Most common H2 topics (deduplicated)
    from collections import Counter
    h2_words = []
    for h in all_h2s:
        if h and h.strip():
            h2_words.extend(h.lower().split()[:4])
    common_topics = [w for w, c in Counter(h2_words).most_common(15) if len(w) > 4 and c > 1]
    
    print(f"  📋 Structure analysis ({n} sites):")
    print(f"     H2 avg: {avg_h2:.1f} | H3 avg: {avg_h3:.1f}")
    print(f"     Tables: {table_pct:.0f}% | Lists: {list_pct:.0f}% | Callouts: {callout_pct:.0f}%")
    print(f"     E-E-A-T signals found: {len(set(experience_phrases))} patterns")
    print(f"     Common H2 topics: {common_topics[:8]}")
    
    # Extract brand/casino names from competitor pages
    import re as _re2
    all_brands = []
    for r in results:
        # Look for casino-like brand names (Title Case, 2-20 chars) near rating/bonus context
        h2_brands = [h for h in r.get("h2", []) if h and len(h.split()) <= 4 and h[0].isupper()]
        all_brands.extend(h2_brands[:3])
    # Deduplicate and filter common words
    STOP_WORDS = {"top", "best", "the", "how", "what", "why", "when", "casino", "online", "review"}
    brand_candidates = []
    seen = set()
    for b in all_brands:
        bw = b.lower().strip()
        words = bw.split()
        if len(words) <= 3 and not any(w in STOP_WORDS for w in words) and bw not in seen:
            seen.add(bw)
            brand_candidates.append(b.strip())
    
    return {
        "sites_analyzed": n,
        "avg_h2_count": round(avg_h2),
        "avg_h3_count": round(avg_h3),
        "table_pct": round(table_pct),
        "list_pct": round(list_pct),
        "callout_pct": round(callout_pct),
        "has_experience_phrases": len(set(experience_phrases)) > 2,
        "top_h2_samples": list(dict.fromkeys(all_h2s))[:6],
        "common_topics": common_topics[:10],
        "competitor_brands": brand_candidates[:10],
    }

async def analyze_competitor_deep(client, urls, keyword=""):
    """
    DEEP analysis of competitor pages — extracts SEO metrics like a real audit tool.
    Returns rich dict with content structure, keyword analysis, readability, links data.
    Used to give the model precise targets to beat competitors.
    """
    import re as _re
    from collections import Counter
    import asyncio as _asyncio

    if not urls:
        return None

    # Stop words to filter from LSI candidates (multilingual)
    STOP_LSI = {
        # EN
        'the','and','for','with','that','this','from','have','will','you','our','are','was','what','how',
        'all','can','any','one','two','also','just','more','than','about','into','only','very','most',
        # NL
        'het','een','van','en','is','dat','in','op','voor','met','zijn','niet','aan','ook','wij','als','door',
        'kan','heb','heeft','zal','dan','maar','wat','om','onze','deze','bij','naar','uit','over','tot',
        # IT
        'il','la','le','di','che','un','una','per','con','del','della','sono','non','si','come','anche','ma',
        'gli','dei','delle','negli','sulla','sul','tra','tra','molto','più','tutti','tutto',
        # DE
        'der','die','das','und','ist','nicht','von','zu','mit','auf','für','eine','einen','sind','als','auch',
        'ein','dem','den','des','im','am','wie','aber','dass','wenn','wir','sie','er','es',
        # PL
        'jest','nie','sie','się','na','do','w','z','i','to','co','oraz','tak','dla','te','jak','tylko',
        'tym','tej','tych','przez','po','od','o','ten','ta','ale','więc','jeszcze','bardzo','wszystkie',
        # FR/ES/PT (basics)
        'les','des','est','sont','dans','sur','plus','tout','tous','aussi','mais','par','vous','nos',
        'el','los','las','para','con','que','en','por','su','este','esta','muy','más','pero',
    }

    async def fetch_deep(url):
        try:
            from smart_fetcher import fetch_html as _smart_fetch
            html, _method = await _smart_fetch(url, timeout=15, use_playwright=False)
            if not html:
                return None
        except Exception:
            return None

        # Strip scripts/styles/svg for clean text extraction
        html_clean = _re.sub(r'<(script|style|noscript|svg)[^>]*>.*?</\1>', '', html, flags=_re.DOTALL|_re.I)
        # Detect main content area — try article/main, fallback to body, fallback to whole html
        body_match = _re.search(r'<(?:article|main)[^>]*>(.*?)</(?:article|main)>', html_clean, _re.DOTALL|_re.I)
        if body_match and len(body_match.group(1)) > 5000:
            body_html = body_match.group(1)
        else:
            body_match2 = _re.search(r'<body[^>]*>(.*?)</body>', html_clean, _re.DOTALL|_re.I)
            body_html = body_match2.group(1) if body_match2 else html_clean
            # Strip header/footer/nav from body
            for tag in ('header', 'footer', 'nav', 'aside'):
                body_html = _re.sub(rf'<{tag}[^>]*>.*?</{tag}>', '', body_html, flags=_re.DOTALL|_re.I)

        # Plain text from body
        body_text = _re.sub(r'<[^>]+>', ' ', body_html)
        body_text = _re.sub(r'\s+', ' ', body_text).strip()
        body_words = body_text.split()
        body_word_count = len(body_words)
        # Total word count (full HTML)
        total_text = _re.sub(r'<[^>]+>', ' ', html_clean)
        total_words_count = len(total_text.split())

        # Headings
        h1s = [_re.sub(r"<[^>]+>", "", h).strip() for h in _re.findall(r"<h1[^>]*>(.*?)</h1>", html_clean, _re.DOTALL|_re.I)]
        h2s = [_re.sub(r"<[^>]+>", "", h).strip() for h in _re.findall(r"<h2[^>]*>(.*?)</h2>", html_clean, _re.DOTALL|_re.I)]
        h3s = [_re.sub(r"<[^>]+>", "", h).strip() for h in _re.findall(r"<h3[^>]*>(.*?)</h3>", html_clean, _re.DOTALL|_re.I)]
        h4s = [_re.sub(r"<[^>]+>", "", h).strip() for h in _re.findall(r"<h4[^>]*>(.*?)</h4>", html_clean, _re.DOTALL|_re.I)]

        # Paragraphs
        para_matches = _re.findall(r'<p[^>]*>(.*?)</p>', body_html, _re.DOTALL|_re.I)
        para_lengths = []
        for p in para_matches:
            p_text = _re.sub(r'<[^>]+>', ' ', p)
            p_words = len(p_text.split())
            if p_words >= 5:
                para_lengths.append(p_words)
        avg_para = sum(para_lengths) / len(para_lengths) if para_lengths else 0
        max_para = max(para_lengths) if para_lengths else 0

        # Tables
        table_matches = _re.findall(r'<table[^>]*>(.*?)</table>', body_html, _re.DOTALL|_re.I)
        table_count = len(table_matches)
        rows_per_table = []
        for tbl in table_matches:
            rows = len(_re.findall(r'<tr', tbl, _re.I))
            if rows: rows_per_table.append(rows)
        avg_table_rows = sum(rows_per_table) / len(rows_per_table) if rows_per_table else 0

        # Lists
        list_matches = _re.findall(r'<(ul|ol)[^>]*>(.*?)</\1>', body_html, _re.DOTALL|_re.I)
        list_count = len(list_matches)
        items_per_list = []
        for tag, content in list_matches:
            items = len(_re.findall(r'<li', content, _re.I))
            if items: items_per_list.append(items)
        avg_list_items = sum(items_per_list) / len(items_per_list) if items_per_list else 0

        # FAQ count (heuristic: count question-like patterns)
        # Detect: <h3>?</h3>, <summary>?</summary>, <dt>?</dt>, schema.org FAQ
        faq_count = (
            len(_re.findall(r'<h[3-4][^>]*>[^<]*\?\s*</h[3-4]>', html_clean, _re.I))
            + len(_re.findall(r'<summary[^>]*>[^<]*\?\s*</summary>', html_clean, _re.I))
            + len(_re.findall(r'<dt[^>]*>[^<]*\?\s*</dt>', html_clean, _re.I))
        )
        # FAQPage schema
        faq_schema_count = len(_re.findall(r'"@type"\s*:\s*"Question"', html_clean))
        faq_count = max(faq_count, faq_schema_count)

        # Images
        img_tags = _re.findall(r'<img[^>]*>', body_html, _re.I)
        images_total = len(img_tags)
        images_with_alt = sum(1 for img in img_tags if _re.search(r'alt\s*=\s*["\']([^"\']+)["\']', img))

        # Schema.org blocks
        schema_blocks = len(_re.findall(r'<script\s+type=["\']application/ld\+json["\']', html_clean, _re.I))

        # Links
        all_links = _re.findall(r'<a\s+[^>]*href\s*=\s*["\']([^"\']+)["\'][^>]*>(.*?)</a>', body_html, _re.DOTALL|_re.I)
        # Determine domain
        url_domain_match = _re.search(r'https?://([^/]+)', url)
        url_domain = url_domain_match.group(1) if url_domain_match else ''
        internal_links = 0
        external_links = 0
        ext_domains = set()
        anchor_words_total = 0
        anchor_with_kw = 0
        for href, anchor_html in all_links:
            anchor_text = _re.sub(r'<[^>]+>', ' ', anchor_html).strip()
            if not anchor_text or anchor_text.lower() in ('home','top','more','read more'):
                continue
            anchor_words_total += len(anchor_text.split())
            if keyword and keyword.lower() in anchor_text.lower():
                anchor_with_kw += 1
            if href.startswith('#') or href.startswith('/') or url_domain in href:
                internal_links += 1
            elif href.startswith('http'):
                external_links += 1
                ext_match = _re.search(r'https?://([^/]+)', href)
                if ext_match:
                    ext_domains.add(ext_match.group(1))

        total_links = internal_links + external_links
        avg_anchor_words = anchor_words_total / total_links if total_links else 0
        links_per_1000 = total_links / max(body_word_count, 1) * 1000

        # Keyword analysis
        kw_count = 0
        kw_in_h1 = False
        kw_in_h2_count = 0
        kw_in_h3_count = 0
        kw_in_first100 = False
        kw_in_last100 = False
        kw_in_anchors = anchor_with_kw
        if keyword:
            kw_lower = keyword.lower()
            body_lower = body_text.lower()
            kw_count = body_lower.count(kw_lower)
            kw_in_h1 = any(kw_lower in h.lower() for h in h1s)
            kw_in_h2_count = sum(1 for h in h2s if kw_lower in h.lower())
            kw_in_h3_count = sum(1 for h in h3s if kw_lower in h.lower())
            first100 = ' '.join(body_words[:100]).lower()
            last100 = ' '.join(body_words[-100:]).lower()
            kw_in_first100 = kw_lower in first100
            kw_in_last100 = kw_lower in last100

        # Img alt KW count
        kw_in_alt = 0
        if keyword:
            for img in img_tags:
                alt_match = _re.search(r'alt\s*=\s*["\']([^"\']+)["\']', img)
                if alt_match and keyword.lower() in alt_match.group(1).lower():
                    kw_in_alt += 1

        # Title and Meta Description
        title_match = _re.search(r'<title[^>]*>(.*?)</title>', html_clean, _re.DOTALL|_re.I)
        title_text = _re.sub(r'\s+', ' ', _re.sub(r'<[^>]+>', '', title_match.group(1))).strip() if title_match else ''
        meta_desc_match = _re.search(r'<meta[^>]*name\s*=\s*["\']description["\'][^>]*content\s*=\s*["\']([^"\']+)["\']', html_clean, _re.I)
        meta_desc_text = meta_desc_match.group(1).strip() if meta_desc_match else ''

        # Slug
        url_path = _re.sub(r'https?://[^/]+', '', url).strip('/')
        slug = url_path.split('/')[-1] if url_path else ''

        # Sentences for readability
        sentences = _re.split(r'(?<=[.!?])\s+', body_text)
        sentences = [s for s in sentences if len(s.split()) >= 3]
        sentence_lengths = [len(s.split()) for s in sentences]
        avg_sentence_len = sum(sentence_lengths) / len(sentence_lengths) if sentence_lengths else 0
        long_sentences = sum(1 for s in sentence_lengths if s > 30)
        long_sent_pct = long_sentences / len(sentence_lengths) * 100 if sentence_lengths else 0

        # LSI: extract most frequent meaningful words (3+ chars, not stopwords, not keyword parts)
        kw_parts = set(keyword.lower().split()) if keyword else set()
        lsi_words = []
        for w in body_words:
            wl = w.lower().strip('.,!?:;()[]{}"\'')
            if (len(wl) >= 4 and wl not in STOP_LSI and wl not in kw_parts
                and not wl.isdigit() and wl.isalpha()):
                lsi_words.append(wl)
        lsi_counter = Counter(lsi_words)
        top_lsi = lsi_counter.most_common(30)
        lsi_unique = len(set(lsi_words))
        lsi_total = len(lsi_words)

        # Long-tail phrases (3-5 word phrases that contain keyword parts)
        long_tail = set()
        if keyword:
            for i in range(len(body_words) - 2):
                # 3-word phrases containing keyword parts
                triple = ' '.join(body_words[i:i+3]).lower().strip('.,!?:;')
                if any(kp in triple for kp in kw_parts) and len(triple) > 15:
                    long_tail.add(triple)

        return {
            'url': url,
            'total_words': total_words_count,
            'body_words': body_word_count,
            'body_ratio': round(body_word_count / max(total_words_count, 1) * 100, 1),
            'avg_paragraph': round(avg_para, 1),
            'max_paragraph': max_para,
            'h1_count': len(h1s),
            'h2_count': len(h2s),
            'h3_count': len(h3s),
            'h4_count': len(h4s),
            'table_count': table_count,
            'avg_table_rows': round(avg_table_rows, 1),
            'list_count': list_count,
            'avg_list_items': round(avg_list_items, 1),
            'faq_count': faq_count,
            'images_total': images_total,
            'images_with_alt': images_with_alt,
            'schema_blocks': schema_blocks,
            'internal_links': internal_links,
            'external_links': external_links,
            'unique_ext_domains': len(ext_domains),
            'ext_domains_sample': list(ext_domains)[:10],
            'avg_anchor_words': round(avg_anchor_words, 1),
            'links_per_1000_words': round(links_per_1000, 1),
            'kw_in_anchors': kw_in_anchors,
            'kw_count': kw_count,
            'kw_density_pct': round(kw_count / max(body_word_count, 1) * 100, 2),
            'kw_in_h1': kw_in_h1,
            'kw_in_h2_count': kw_in_h2_count,
            'kw_in_h3_count': kw_in_h3_count,
            'kw_in_first100': kw_in_first100,
            'kw_in_last100': kw_in_last100,
            'kw_in_alt': kw_in_alt,
            'title': title_text,
            'title_length': len(title_text),
            'meta_desc': meta_desc_text,
            'meta_desc_length': len(meta_desc_text),
            'slug': slug,
            'avg_sentence_len': round(avg_sentence_len, 1),
            'long_sentence_pct': round(long_sent_pct, 1),
            'top_lsi': top_lsi[:20],
            'lsi_unique': lsi_unique,
            'lsi_total': lsi_total,
            'long_tail_count': len(long_tail),
            'long_tail_sample': list(long_tail)[:15],
            'h1_sample': h1s[:1],
            'h2_samples': h2s[:8],
        }

    tasks = [fetch_deep(url) for url in urls[:8]]
    raw = await _asyncio.gather(*tasks)
    results = [r for r in raw if r]

    if not results:
        return None

    n = len(results)

    # Aggregate
    def avg(field):
        vals = [r.get(field, 0) for r in results if r.get(field) is not None]
        return round(sum(vals) / len(vals), 1) if vals else 0

    # Aggregate top LSI from all sites
    all_lsi = Counter()
    for r in results:
        for w, c in r.get('top_lsi', []):
            all_lsi[w] += c
    top_lsi_global = all_lsi.most_common(20)

    # Common ext domains across competitors
    domain_counter = Counter()
    for r in results:
        for d in r.get('ext_domains_sample', []):
            domain_counter[d] += 1
    common_domains = [d for d, c in domain_counter.most_common(15) if c >= 2]

    # Aggregate long-tail samples
    all_long_tail = []
    for r in results:
        all_long_tail.extend(r.get('long_tail_sample', []))

    summary = {
        'sites_analyzed': n,
        'urls': [r['url'] for r in results],
        # Content structure (averaged)
        'avg_total_words': avg('total_words'),
        'avg_body_words': avg('body_words'),
        'avg_body_ratio': avg('body_ratio'),
        'avg_paragraph_words': avg('avg_paragraph'),
        'max_paragraph_words': max((r.get('max_paragraph', 0) for r in results), default=0),
        'avg_h2_count': avg('h2_count'),
        'avg_h3_count': avg('h3_count'),
        'avg_h4_count': avg('h4_count'),
        'avg_table_count': avg('table_count'),
        'avg_table_rows': avg('avg_table_rows'),
        'avg_list_count': avg('list_count'),
        'avg_list_items': avg('avg_list_items'),
        'avg_faq_count': avg('faq_count'),
        'max_faq_count': max((r.get('faq_count', 0) for r in results), default=0),
        'avg_images': avg('images_total'),
        'avg_images_alt_pct': round(sum((r.get('images_with_alt', 0) / max(r.get('images_total', 1), 1) * 100) for r in results) / n, 1),
        'avg_schema_blocks': avg('schema_blocks'),
        # Links (averaged)
        'avg_internal_links': avg('internal_links'),
        'avg_external_links': avg('external_links'),
        'avg_unique_ext_domains': avg('unique_ext_domains'),
        'avg_anchor_words': avg('avg_anchor_words'),
        'avg_links_per_1000_words': avg('links_per_1000_words'),
        'common_external_domains': common_domains,
        # Keyword (averaged)
        'avg_kw_count': avg('kw_count'),
        'avg_kw_density_pct': avg('kw_density_pct'),
        'kw_in_h1_pct': round(sum(1 for r in results if r.get('kw_in_h1')) / n * 100, 0),
        'avg_kw_in_h2': avg('kw_in_h2_count'),
        'avg_kw_in_h3': avg('kw_in_h3_count'),
        'kw_in_first100_pct': round(sum(1 for r in results if r.get('kw_in_first100')) / n * 100, 0),
        'kw_in_last100_pct': round(sum(1 for r in results if r.get('kw_in_last100')) / n * 100, 0),
        'avg_kw_in_anchors': avg('kw_in_anchors'),
        'avg_kw_in_alt': avg('kw_in_alt'),
        # Readability (averaged)
        'avg_sentence_len': avg('avg_sentence_len'),
        'avg_long_sentence_pct': avg('long_sentence_pct'),
        # On-page SEO (averaged)
        'avg_title_length': avg('title_length'),
        'avg_meta_desc_length': avg('meta_desc_length'),
        'sample_titles': [r['title'] for r in results if r.get('title')][:3],
        # LSI semantic terms
        'avg_lsi_unique': avg('lsi_unique'),
        'avg_lsi_total': avg('lsi_total'),
        'top_lsi_terms': top_lsi_global[:20],
        # Long-tail
        'avg_long_tail_count': avg('long_tail_count'),
        'long_tail_samples': list(set(all_long_tail))[:20],
        # H1/H2 examples
        'h1_samples': [r.get('h1_sample', [''])[0] for r in results if r.get('h1_sample')][:3],
        'h2_samples': list(dict.fromkeys([h for r in results for h in r.get('h2_samples', [])]))[:10],
    }

    # Print compact summary
    print(f"  🔬 Deep competitor analysis ({n} sites):")
    print(f"     Body: {summary['avg_body_words']}w avg | paragraphs avg {summary['avg_paragraph_words']}w | sentences avg {summary['avg_sentence_len']}w")
    print(f"     H2={summary['avg_h2_count']} | H3={summary['avg_h3_count']} | Tables={summary['avg_table_count']} ({summary['avg_table_rows']} rows) | Lists={summary['avg_list_count']} | FAQ={summary['avg_faq_count']}")
    print(f"     Images={summary['avg_images']} ({summary['avg_images_alt_pct']}% alt) | Schema={summary['avg_schema_blocks']}")
    print(f"     Links: int={summary['avg_internal_links']}, ext={summary['avg_external_links']} ({summary['avg_unique_ext_domains']} unique domains, {summary['avg_links_per_1000_words']}/1000w)")
    if keyword:
        print(f"     KW '{keyword}': {summary['avg_kw_count']}x ({summary['avg_kw_density_pct']}%) | in H1: {summary['kw_in_h1_pct']:.0f}% sites | H2: {summary['avg_kw_in_h2']}x | first100: {summary['kw_in_first100_pct']:.0f}%")
    print(f"     LSI: {summary['avg_lsi_unique']} unique terms (top: {[w for w,c in summary['top_lsi_terms'][:10]]})")
    print(f"     Long-tail: {summary['avg_long_tail_count']} phrases")
    if common_domains:
        print(f"     Common ext links: {common_domains[:5]}")

    return summary


async def analyze_competitor_tf(client, urls, keyword):
    """
    Analyze TF (Term Frequency) of keyword across competitor pages.
    Returns dict with avg_tf, avg_count, avg_words, target_count for our text.
    
    TF = keyword_count / total_words
    Target: match avg competitor TF ± 10%
    """
    import re as _re
    
    if not urls or not keyword:
        return None
    
    results = []
    kw_lower = keyword.lower()
    
    async def fetch_tf(url):
        for attempt in range(1):
            try:
                from smart_fetcher import fetch_html as _smart_fetch
                _html, _method = await _smart_fetch(url, timeout=12, use_playwright=False)
                if _html:
                    # Strip HTML
                    text = _re.sub(r"<[^>]+>", " ", _html)
                    text = _re.sub(r"\s+", " ", text).strip().lower()
                    words = text.split()
                    total = len(words)
                    if total < 200:
                        return None  # too short, skip
                    count = len(_re.findall(_re.escape(kw_lower), text))
                    tf = count / total if total else 0
                    tf_per_100 = tf * 100
                    return {"url": url, "count": count, "total": total, "tf": tf, "tf_per_100": tf_per_100}
            except Exception:
                pass
        return None
    
    import asyncio
    tasks = [fetch_tf(url) for url in urls[:10]]
    raw = await asyncio.gather(*tasks)
    results = [r for r in raw if r]
    
    if len(results) < 3:
        return None  # not enough data
    
    avg_tf_per_100 = sum(r["tf_per_100"] for r in results) / len(results)
    avg_count = sum(r["count"] for r in results) / len(results)
    avg_words = sum(r["total"] for r in results) / len(results)
    
    print(f"  📊 TF analysis ({len(results)} sites): avg TF={avg_tf_per_100:.2f}% | avg count={avg_count:.1f}x / {avg_words:.0f} words")
    for r in results[:5]:
        print(f"    {r['url'][:50]}: {r['count']}x / {r['total']} words = TF {r['tf_per_100']:.2f}%")
    
    return {
        "keyword": keyword,
        "sites_analyzed": len(results),
        "avg_tf_per_100": round(avg_tf_per_100, 2),
        "avg_count": round(avg_count, 1),
        "avg_words": round(avg_words),
        "target_tf_min": round(avg_tf_per_100 * 0.8, 2),   # -20% of avg
        "target_tf_max": round(avg_tf_per_100 * 1.2, 2),   # +20% of avg
        "recommendation": f"Use keyword {round(avg_count * 0.9)}-{round(avg_count * 1.1)} times in ~{round(avg_words)} words (TF: {avg_tf_per_100:.2f}%)"
    }

async def get_competitor_wordcount(client, urls):
    """
    Analyse competitor word counts with graceful degradation:
    - Try all provided URLs (up to 10), timeout 10s + 1 retry each
    - If >= 3 accessible: compute average + 20%, proceed
    - If < 3 accessible: return None → caller asks user to replace competitors
    - Ignore pages with < 300 words (likely error pages / paywalls)
    """
    if not urls:
        return 2000  # no competitors provided — use default

    import re as _re
    MIN_ACCESSIBLE = 3
    counts = []
    failed = []

    async def fetch_wordcount(url):
        """Fetch via smart_fetcher (curl_cffi → httpx → playwright)"""
        for attempt in range(1):
            try:
                from smart_fetcher import fetch_html as _smart_fetch
                _html, _method = await _smart_fetch(url, timeout=12, use_playwright=False)
                if _html:
                    text = _re.sub(r"<[^>]+>", " ", _html)
                    text = _re.sub(r"\s+", " ", text).strip()
                    words = len(text.split())
                    if words > 300:
                        return words
                    return None  # page too short — skip
                # Non-200 on attempt 0 → retry; on attempt 1 → fail
            except Exception:
                pass  # retry on exception
        return None  # both attempts failed

    import asyncio
    tasks = [fetch_wordcount(url) for url in urls[:10]]
    results = await asyncio.gather(*tasks, return_exceptions=False)

    for url, wc in zip(urls[:10], results):
        short_url = url[:55]
        if wc:
            counts.append(wc)
            print(f"  ✅ {short_url}: {wc} words")
        else:
            failed.append(url)
            print(f"  ❌ {short_url}: unavailable")

    accessible = len(counts)
    total = min(len(urls), 10)
    print(f"  Accessible: {accessible}/{total}")

    if accessible < MIN_ACCESSIBLE:
        print(f"  ⚠️  Only {accessible} accessible — need at least {MIN_ACCESSIBLE}")
        return None  # signal to caller: ask user to replace competitors

    avg = int(sum(counts) / accessible)
    result = int(avg * 1.2)           # +20% vs TOP competitors
    result = max(1500, min(result, 4000))  # clamp 1500–7000
    print(f"  Avg({accessible} sites): {avg} words → +20%: {result} words")
    return result


async def check_originality(text: str, api_key: str = "e36aslfpx1vy49muckzn2orh0t7gi8wb") -> dict:
    """Check text originality via Originality.ai API."""
    import httpx
    sample = text[:3000] if len(text) > 3000 else text
    try:
        async with httpx.AsyncClient(timeout=30) as client:
            r = await client.post(
                "https://api.originality.ai/api/v1/scan/ai",
                headers={"X-OAI-API-KEY": api_key, "Content-Type": "application/json"},
                json={"content": sample, "title": "content_check"}
            )
            if r.status_code == 200:
                data = r.json()
                score = data.get("score", {})
                original_pct = round(score.get("original", 0) * 100, 1)
                ai_pct = round(score.get("ai", 0) * 100, 1)
                return {"original": original_pct, "ai": ai_pct, "ok": original_pct >= 50}
    except Exception as e:
        print(f"  ⚠️ Originality.ai error: {e}")
    return {"original": 0, "ai": 100, "ok": False}


def _health_check():
    """Check OR key and Gemini key before starting generation."""
    import urllib.request as _ur, json as _jh
    print("\n🔍 Health check...")
    try:
        req = _ur.Request(
            "https://openrouter.ai/api/v1/auth/key",
            headers={"Authorization": f"Bearer {OPENROUTER_KEY}"}
        )
        resp = _jh.loads(_ur.urlopen(req, timeout=8).read())
        remaining = resp.get("data", {}).get("limit_remaining", 0)
        print(f"  ✅ OR key OK | remaining: ${remaining:.1f}")
    except Exception as e:
        raise RuntimeError(f"OpenRouter key check failed: {e}")
    if GEMINI_KEY:
        print(f"  ✅ Gemini key present (used via OpenRouter, no direct check needed)")
    print("  ✅ Health check passed\n")

async def main():
    with open(TASK_PATH, encoding="utf-8") as f:
        task = json.load(f)
    
    # AUTO-FIX disabled — count is always taken from task.json only
    
    casinos = task.get("casinos", [])
    count = task.get("count", 5)

    # Determine output format: html_txt or docx (default)
    # Also check format.txt file in workspace (agent can write this)
    _fmt_file = os.path.join(WORKSPACE, "format.txt")
    _fmt_from_file = ""
    if os.path.exists(_fmt_file):
        try:
            _fmt_from_file = open(_fmt_file).read().strip().lower()
            os.remove(_fmt_file)  # auto-delete after reading
        except: pass
    output_format = task.get("output_format", _fmt_from_file or "docx").lower()
    # Also detect from requirements/topic/notes fields
    # Check all text fields for html_txt hints — search the entire task.json as string
    import json as _json
    _all_fields = _json.dumps(task, ensure_ascii=False).lower()
    _html_keywords = ["для сайта", "для постинга", "на сайт", "пост на сайт",
                     "html", "for site", "for posting", "website", "html tags", "html тег",
                     "html_txt", "txt format", "формат html", "формат txt"]
    if output_format != "html_txt" and any(kw in _all_fields.lower() for kw in _html_keywords):
        output_format = "html_txt"
    use_html_txt = output_format == "html_txt"
    _ext = "txt" if use_html_txt else "docx"
    if use_html_txt:
        print(f"  📄 Output format: HTML TXT (for website posting)")
    # Check tz_count.txt override — deleted after reading to avoid stale values
    tz_count_file = os.path.join(os.path.dirname(TASK_PATH), "tz_count.txt")
    if os.path.exists(tz_count_file):
        try:
            fc = int(open(tz_count_file).read().strip())
            if fc != count:
                print(f"  🔧 count override from tz_count.txt: {count} → {fc}")
                count = fc
        except: pass
        finally:
            try: os.remove(tz_count_file)
            except: pass
    date = datetime.datetime.now().strftime("%Y-%m-%d_%H%M")
    geo_slug = task.get("geo", "uk").lower().replace(" ", "-")
    word_count = task.get("word_count") # Инициализация word_count
    
    print(f"Задача: {count} текстов, {task.get('geo')}, {task.get('language')}")
    print(f"Казино: {casinos}")
    print(f"Gemini key: {GEMINI_KEY[:20]}..." if GEMINI_KEY else "❌ Нет Gemini ключа!")
    
    if not GEMINI_KEY:
        print("ОШИБКА: Не найден Gemini API ключ в SECRETS.md")
        return
    
    # Если word_count не задан или None — анализируем конкурентов
    if not word_count or word_count is None or (word_count == 4000 and task.get("word_count") is None):
        competitors = task.get("competitors", [])
        if competitors:
            print(f"\nАнализируем конкурентов для определения объёма...")
            async with httpx.AsyncClient() as tmp_client:
                word_count = await get_competitor_wordcount(tmp_client, competitors)
                # Analyze TF-IDF for primary keyword
                primary_kw = task.get("keywords", [""])[0] if task.get("keywords") else ""
                if primary_kw and competitors:
                    tf_data = await analyze_competitor_tf(tmp_client, competitors, primary_kw)
                    if tf_data:
                        task["competitor_tf"] = tf_data
                        print(f"  ✅ TF target: {tf_data['recommendation']}")
                # Deep structural analysis of competitors
                struct_data = await analyze_competitor_structure(tmp_client, competitors, primary_kw)
                if struct_data:
                    task["competitor_structure"] = struct_data
                    if struct_data.get("competitor_brands") and not task.get("casinos"):
                        task["competitor_brands"] = struct_data["competitor_brands"]
                        print(f"  🏷️ Brands from competitors: {struct_data['competitor_brands'][:5]}")
                    print(f"  ✅ Structure: {struct_data['avg_h2_count']} H2s, tables={struct_data['table_pct']}%, lists={struct_data['list_pct']}%")

                # FULL DEEP analysis (LSI, links, paragraphs, FAQ, images, on-page SEO)
                deep_data = await analyze_competitor_deep(tmp_client, competitors, primary_kw)
                if deep_data:
                    task["competitor_deep"] = deep_data

            if word_count is None:
                # Less than 3 competitors accessible — use fallback word count instead of aborting
                print(f"  ⚠️ Less than 3 competitors accessible — using fallback word count 5000")
                word_count = 5000

            word_count = int(word_count)
            print(f"Целевой объём: {word_count} слов")
        else:
            word_count = 2000
            print(f"Конкуренты не указаны, объём по умолчанию: {word_count} слов")
    task["word_count"] = word_count  # persist for build_prompt

    # ── v3: Compute dynamic SERP targets ─────────────────────────────────────
    serp_targets = None
    primary_kw = task.get("keywords", [""])[0] if task.get("keywords") else ""
    if primary_kw:
        print("\n  🎯 v3: Computing SERP targets via Ahrefs + scraping...")
        async with httpx.AsyncClient() as _serp_client:
            serp_targets = await get_targets_from_serp(primary_kw, task.get("geo", "US"), _serp_client)
            # Also use SERP URLs as additional competitors if we got them
            serp_urls = serp_targets.pop("_serp_urls", [])
            serp_metrics = serp_targets.pop("_metrics_raw", [])
            if serp_urls and not task.get("competitors"):
                task["competitors"] = serp_urls[:5]
                print(f"  ℹ️  Using {len(task['competitors'])} Ahrefs SERP URLs as competitors")
        task["serp_targets"] = serp_targets
    else:
        serp_targets = _default_targets()

    # Строим промпты
    prompts = []
    filenames = []
    task_type = task.get("article_type", "text_review")
    
    if task_type == "text_mono":
        # text_mono: один текст на каждый бренд
        for casino in casinos[:count]:
            slug = casino.lower().replace(" ", "_").replace("-", "_").replace(".", "")
            prompts.append((task_type, casino))
            filenames.append(f"{slug}_{geo_slug}_mono_{date}.{_ext}")
        # Если брендов меньше count — повторяем с разными фокусами
        while len(prompts) < count:
            casino = casinos[len(prompts) % len(casinos)] if casinos else "casino"
            slug = casino.lower().replace(" ", "_").replace("-", "_").replace(".", "")
            prompts.append((task_type, casino))
            filenames.append(f"{slug}_{geo_slug}_mono_{date}_{len(prompts)}.{_ext}")
    else:
        # text_review: все тексты — многобрендовые обзоры (все бренды в каждом тексте)
        for i in range(count):
            prompts.append((task_type, None))
            filenames.append(f"best_online_casinos_{geo_slug}_{date}_{i+1}.{_ext}")
    
    # Получаем Dropbox токен
    print("Получаем Dropbox токен...")
    try:
        dbx_token = get_dropbox_token()
        print("✅ Dropbox подключён")
    except Exception as e:
        print(f"❌ Dropbox ошибка: {e}")
        dbx_token = None

    # Пре-генерация уникальных titles для всего батча
    print(f"\n🏷️ Пре-генерация {len(prompts)} уникальных titles...")
    pre_titles = []
    try:
        pre_titles = await _pregen_unique_titles(task, len(prompts), serp_targets=serp_targets)
    except Exception as _pt_err:
        print(f"  ⚠️ Pre-title gen error: {_pt_err}")
        pre_titles = []

    # Параллельная генерация
    print(f"\nЗапускаем {len(prompts)} параллельных запросов к Gemini...")
    
    async with httpx.AsyncClient() as client:
        tasks_gen = []
        for i, (atype, casino) in enumerate(prompts):
            pre_title = pre_titles[i] if i < len(pre_titles) else None
            prompt = build_prompt(task, atype, casino, text_index=i, serp_targets=serp_targets, pre_title=pre_title)
            tasks_gen.append(generate_text(client, prompt, i, task=task, serp_targets=serp_targets))
        results = list(await asyncio.gather(*tasks_gen))
    
    # Сохраняем и загружаем
    links = []
    completed = 0
    
    for i, (text, fname) in enumerate(zip(results, filenames)):
        if not text:
            print(f"  [{i+1}] Пропускаем — нет текста")
            continue
        
        local_path = f"/tmp/{fname}"
        if use_html_txt:
            text_to_html_txt(text, local_path)
            print(f"  [{i+1}] HTML TXT: {local_path}")
        else:
            text_to_docx(text, local_path)
            print(f"  [{i+1}] DOCX: {local_path}")
        
        # Originality.ai — blocking check with regeneration on low score (V4)
        ORIG_THRESHOLD = 80  # % minimum originality
        ORIG_MAX_RETRIES = 2
        upload_fname = fname
        orig_score = None

        for _orig_attempt in range(ORIG_MAX_RETRIES + 1):
            try:
                print(f"  [{i+1}] Originality.ai проверка (попытка {_orig_attempt+1}/{ORIG_MAX_RETRIES+1})...")
                _orig_data = asyncio.get_event_loop().run_until_complete(check_originality(text))
                orig_score = _orig_data.get("original", 0)
                print(f"  [{i+1}] Originality: {orig_score}% original (AI: {_orig_data.get('ai', '?')}%)")
            except Exception as _oe:
                print(f"  [{i+1}] Originality check error: {_oe} — пропускаем блокировку")
                orig_score = None
                break

            if orig_score is None or orig_score >= ORIG_THRESHOLD:
                break  # OK

            if _orig_attempt < ORIG_MAX_RETRIES:
                print(f"  [{i+1}] ⚠️  Originality {orig_score}% < {ORIG_THRESHOLD}% — регенерируем текст...")
                _atype, _casino = prompts[i]
                _prompt = build_prompt(task, _atype, _casino, text_index=i, serp_targets=serp_targets)
                async def _regen(_p):
                    async with httpx.AsyncClient() as _c:
                        return await generate_text(_c, _p, i, task=task, serp_targets=serp_targets)
                text = asyncio.get_event_loop().run_until_complete(_regen(_prompt))
                if not text:
                    print(f"  [{i+1}] ❌ Регенерация вернула пустой текст — используем предыдущий")
                    break
                # Пересохраняем файл с новым текстом
                if use_html_txt:
                    text_to_html_txt(text, local_path)
                else:
                    text_to_docx(text, local_path)
            else:
                # Все попытки исчерпаны — загружаем с пометкой
                print(f"  [{i+1}] ⚠️  Originality {orig_score}% после {ORIG_MAX_RETRIES} регенераций — загружаем с пометкой _lowscore")
                _base, _ext_part = fname.rsplit(".", 1) if "." in fname else (fname, "")
                upload_fname = f"{_base}_lowscore.{_ext_part}" if _ext_part else f"{_base}_lowscore"
                local_lowscore = f"/tmp/{upload_fname}"
                import shutil as _sh
                _sh.copy(local_path, local_lowscore)
                local_path = local_lowscore
        
        if dbx_token:
            try:
                dbx_folder = task.get('dropbox_path') or task.get('dropbox_folder', '/sofiya/')
                dbx_path = f"{dbx_folder.rstrip('/')}/{upload_fname}"
                link = upload_to_dropbox(dbx_token, local_path, dbx_path)
                links.append(link)
                print(f"  [{i+1}] ✅ Dropbox: {link}")
                completed += 1
            except Exception as e:
                print(f"  [{i+1}] ❌ Dropbox ошибка: {e}")
                links.append(f"local:/tmp/{fname}")
                completed += 1
        else:
            links.append(f"local:/tmp/{fname}")
            completed += 1
    
    # Обновляем task.json
    task["completed"] = completed
    task["links"] = links
    task["status"] = "done" if completed >= count else "partial"
    
    with open(TASK_PATH, "w", encoding="utf-8") as f:
        json.dump(task, f, indent=2, ensure_ascii=False)
    
    # Обновляем handoff.md
    handoff = f"""# handoff.md — Передача задачи

## Статус: {'done' if completed >= count else 'partial'}

**Кто выполнил:** generate_parallel.py
**Дата:** {date}
**Поставил:** {task.get('assigned_by', '?')}

## Что сделано
- Сгенерировано {completed} из {count} текстов
- Язык: {task.get('language')}
- Тема: {task.get('topic')}
- Казино: {", ".join(casinos)}

## Файлы
"""
    for i, (fname, link) in enumerate(zip(filenames[:completed], links)):
        handoff += f"- [{fname}]({link})\n"
    
    with open(os.path.join(WORKSPACE, "handoff.md"), "w", encoding="utf-8") as f:
        f.write(handoff)
    
    print(f"\n✅ Готово! {completed}/{count} текстов")
    print("Ссылки:")
    for link in links:
        print(f"  - {link}")


    # Записываем result.txt для бота
    result_lines = ["Задача выполнена. {}/{} текстов готовы:".format(completed, count), ""]
    for i, (fname, link) in enumerate(zip(filenames[:completed], links)):
        result_lines.append("{}. {}".format(i+1, link))
    result_text = "\n".join(result_lines)
    with open(os.environ.get("RESULT_FILE_OVERRIDE") or os.path.join(WORKSPACE, "result.txt"), "w", encoding="utf-8") as rf:
        rf.write(result_text)
    print("result.txt записан")
    
    # Отправляем в группу с Софией
    # Send to group with retry and chunking (Telegram limit: 4096 chars)
    MAX_TG = 4000
    parts = []
    if len(result_text) > MAX_TG:
        lines = result_text.split('\n')
        current = ""
        for line in lines:
            if len(current) + len(line) + 1 > MAX_TG:
                parts.append(current)
                current = line
            else:
                current = current + '\n' + line if current else line
        if current:
            parts.append(current)
    else:
        parts = [result_text]

    import time as _time
    for part_idx, part in enumerate(parts):
        for _retry in range(3):
            try:
                import requests as _rq
                _r = _rq.post(
                    f"https://api.telegram.org/bot{TELEGRAM_BOT_TOKEN}/sendMessage",
                    json={"chat_id": TELEGRAM_GROUP_ID, "text": part},
                    timeout=15)
                if _r.status_code == 200:
                    print(f"  ✅ Telegram sent to {TELEGRAM_GROUP_ID} (part {part_idx+1}/{len(parts)})")
                    break
                else:
                    print(f"  ❌ Telegram error {_r.status_code}, retry {_retry+1}")
                    _time.sleep(3)
            except Exception as _te:
                print(f"  ❌ Telegram exception: {_te}, retry {_retry+1}")
                _time.sleep(3)
        if len(parts) > 1:
            _time.sleep(1)

if __name__ == "__main__":
    import sys as _sys2, datetime as _dt2
    _LOG_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "logs")
    os.makedirs(_LOG_DIR, exist_ok=True)
    _log_file = os.path.join(_LOG_DIR, f"generation_{_dt2.datetime.now().strftime('%Y%m%d_%H%M%S')}.log")
    class _Tee:
        def __init__(self, *files):
            self.files = files
        def write(self, data):
            for f in self.files:
                try: f.write(data); f.flush()
                except: pass
        def flush(self):
            for f in self.files:
                try: f.flush()
                except: pass
    with open(_log_file, "w", encoding="utf-8") as _lf:
        _orig_out, _orig_err = _sys2.stdout, _sys2.stderr
        _sys2.stdout = _Tee(_orig_out, _lf)
        _sys2.stderr = _Tee(_orig_err, _lf)
        try:
            print(f"=== Generation started: {_dt2.datetime.now().strftime('%Y-%m-%d %H:%M:%S')} ===")
            print(f"=== Log: {_log_file} ===")
            _health_check()
            asyncio.run(main())
            print(f"=== Generation finished OK: {_dt2.datetime.now().strftime('%Y-%m-%d %H:%M:%S')} ===")
        except Exception as _e:
            import traceback as _tb
            print(f"\n=== FATAL ERROR: {_e} ===")
            _tb.print_exc()
            try:
                import requests as _req2
                _req2.post(
                    f"https://api.telegram.org/bot{TELEGRAM_BOT_TOKEN}/sendMessage",
                    json={"chat_id": TELEGRAM_GROUP_ID, "text": f"\u274c \u0413\u0435\u043d\u0435\u0440\u0430\u0446\u0438\u044f \u0437\u0430\u0432\u0435\u0440\u0448\u0438\u043b\u0430\u0441\u044c \u0441 \u043e\u0448\u0438\u0431\u043a\u043e\u0439:\n{_e}"},
                    timeout=10
                )
            except: pass
        finally:
            _sys2.stdout = _orig_out
            _sys2.stderr = _orig_err
    print(f"\n📋 Log saved: {_log_file}")